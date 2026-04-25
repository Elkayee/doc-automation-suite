"""
Script chuyen doi .md sang .docx
Co ho tro: Heading, Bold, Italic, Tables, Code blocks, Bullet lists
Mermaid:   Render qua mermaid.ink API => PNG => chen vao DOCX
Usage: python convert_md_to_docx.py [input.md] [output.docx]
"""
import re
import base64
import io
import os
import sys
import requests
import time

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_BASE = r"c:\Users\Home33\IdeaProjects\NMCNPM"
_default_md   = os.path.join(_BASE, "Bao_Cao_Tieu_Luan_NMCNPM.md")
_default_docx = os.path.join(_BASE, "Bao_Cao_Tieu_Luan_NMCNPM.docx")

MD_FILE   = sys.argv[1] if len(sys.argv) > 1 else _default_md
DOCX_FILE = sys.argv[2] if len(sys.argv) > 2 else MD_FILE.replace(".md", ".docx")
IMG_CACHE = os.path.join(os.path.dirname(MD_FILE), "mermaid_cache")

os.makedirs(IMG_CACHE, exist_ok=True)

# === MAU SAC ================================================================
COLOR_H1 = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_H2 = RGBColor(0x1F, 0x61, 0x9E)
COLOR_H3 = RGBColor(0x2E, 0x86, 0xAB)
COLOR_H4 = RGBColor(0x44, 0x9D, 0xD1)

# === MERMAID RENDER =========================================================

import struct, zlib

def get_png_dimensions(path: str) -> tuple[int, int]:
    """Doc width, height tu PNG header (bytes 16-24) ma khong can PIL."""
    with open(path, 'rb') as f:
        f.read(8)   # PNG signature
        f.read(4)   # chunk length
        f.read(4)   # 'IHDR'
        w = struct.unpack('>I', f.read(4))[0]
        h = struct.unpack('>I', f.read(4))[0]
    return w, h


def render_mermaid(code: str, diagram_index: int) -> str | None:
    """
    Gui code Mermaid len mermaid.ink (tu dong chon kich thuoc),
    luu PNG vao cache. Tra ve duong dan file hoac None neu loi.
    """
    cache_file = os.path.join(IMG_CACHE, f"diagram_{diagram_index:03d}.png")

    if os.path.exists(cache_file):
        print(f"  [cache] diagram_{diagram_index:03d}.png")
        return cache_file

    encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
    # type=png bat buoc tra PNG; khong dat width/height → tu co theo noi dung bieu do
    url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"

    try:
        print(f"  [render] Dang render diagram {diagram_index}...")
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200 and resp.headers.get('content-type', '').startswith('image'):
            with open(cache_file, 'wb') as f:
                f.write(resp.content)
            print(f"  [OK] Luu vao {cache_file}")
            time.sleep(0.5)  # Tranh spam API
            return cache_file
        else:
            print(f"  [WARN] API tra ve {resp.status_code} cho diagram {diagram_index}")
            return None
    except Exception as e:
        print(f"  [ERROR] Khong the render diagram {diagram_index}: {e}")
        return None


# === DOCX HELPERS ===========================================================

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_para_shading(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)


def set_page_setup(doc: Document):
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)


def strip_md_links(text: str) -> str:
    """Chuyen [text](url) va [text](#anchor) thanh chi giu text."""
    return re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)


def add_formatted_run(para, text: str):
    """Phan tach text thanh run: **bold**, *italic*, `code`, [link](#)."""
    # Xu ly link truoc: [text](url) → text
    text = strip_md_links(text)
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for token in tokens:
        if not token:
            continue
        run = para.add_run()
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name  = 'Courier New'
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run.text = token


def strip_md_markup(text: str) -> str:
    """Xoa **bold**, *italic*, `code` markup de dung o noi khong can run."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*',     r'\1', text)
    text = re.sub(r'`([^`]+)`',       r'\1', text)
    return text


# === MAIN PARSER ============================================================

def parse_and_write(doc: Document, md_path: str):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    # Font mac dinh
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    i              = 0
    in_code        = False
    mermaid_block  = False
    mermaid_buf    = []
    diagram_idx    = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Code/Mermaid block bat dau hoac ket thuc ──────────────────────
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                lang    = line.strip()[3:].strip().lower()
                mermaid_block = (lang == 'mermaid')
                mermaid_buf   = []
            else:
                # Dong code block
                if mermaid_block and mermaid_buf:
                    diagram_idx += 1
                    code = '\n'.join(mermaid_buf)
                    img_path = render_mermaid(code, diagram_idx)
                    if img_path:
                        # Chen anh vao doc
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        # Doc kich thuoc that cua anh PNG
                        img_w, img_h = get_png_dimensions(img_path)
                        # PAGE_W = 6.2 inch (net trang), PAGE_H = 8.0 inch (chieu cao kha dung)
                        PAGE_W = Inches(6.2)
                        PAGE_H = Inches(8.0)
                        # Tinh ti le khung hinh
                        aspect = img_h / img_w if img_w > 0 else 1.0
                        if aspect > 1.2:
                            # Bieu do DOC (flowchart TD, sequence, v.v.)
                            # → gioi han theo chieu cao de vua trang, cho phep rong nho hon
                            calc_h = PAGE_H
                            calc_w = Inches(min(6.2, PAGE_H.inches / aspect))
                            run.add_picture(img_path, width=calc_w, height=calc_h)
                        else:
                            # Bieu do NGANG (graph LR, erDiagram, v.v.) → full width
                            run.add_picture(img_path, width=PAGE_W)
                        # Caption nho
                        cap = doc.add_paragraph(f"Bieu do {diagram_idx}")
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.runs[0].font.size    = Pt(10)
                        cap.runs[0].italic       = True
                        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    else:
                        # Fallback: ghi ro placeholder
                        p = doc.add_paragraph()
                        r = p.add_run(f"[Bieu do Mermaid {diagram_idx} — khong the render]")
                        r.italic = True
                        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                elif not mermaid_block:
                    # Ket thuc code block binh thuong: them dong trong de phan cach
                    pass
                in_code       = False
                mermaid_block = False
                mermaid_buf   = []
            i += 1
            continue

        if in_code:
            if mermaid_block:
                mermaid_buf.append(line)
            else:
                p   = doc.add_paragraph()
                run = p.add_run(line if line else ' ')
                run.font.name      = 'Courier New'
                run.font.size      = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                set_para_shading(p, 'F4F4F4')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
            i += 1
            continue

        # ── Bang (Table) ───────────────────────────────────────────────────
        if line.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                table_rows.append(cells)
                i += 1
            # Bo dong separator |---|
            data_rows = [r for r in table_rows
                         if not all(re.match(r'^[-: ]+$', c) for c in r)]
            if not data_rows:
                continue
            max_cols  = max(len(r) for r in data_rows)
            data_rows = [r + [''] * (max_cols - len(r)) for r in data_rows]

            tbl = doc.add_table(rows=len(data_rows), cols=max_cols)
            tbl.style = 'Table Grid'
            for ri, row in enumerate(data_rows):
                for ci, cell_text in enumerate(row):
                    cell  = tbl.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    p     = cell.paragraphs[0]
                    clean = strip_md_markup(cell_text)
                    run   = p.add_run(clean)
                    run.font.size = Pt(11)
                    if ri == 0:
                        run.bold           = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_cell_bg(cell, '1F619E')
                    elif ri % 2 == 0:
                        set_cell_bg(cell, 'EBF4FB')
            doc.add_paragraph()
            continue

        # ── Horizontal rule ────────────────────────────────────────────────
        if re.match(r'^---+\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # ── Heading ────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level    = len(m.group(1))
            text     = strip_md_markup(m.group(2))
            h_style  = f'Heading {min(level, 4)}'
            p        = doc.add_paragraph(style=h_style)
            run      = p.add_run(text)
            run.font.name = 'Times New Roman'
            colors   = [COLOR_H1, COLOR_H2, COLOR_H3, COLOR_H4, COLOR_H4, COLOR_H4]
            sizes    = [16, 14, 13, 12, 12, 11]
            run.font.color.rgb = colors[level - 1]
            run.font.size      = Pt(sizes[level - 1])
            run.bold           = (level <= 3)
            i += 1
            continue

        # ── Blockquote ─────────────────────────────────────────────────────
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            p    = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run  = p.add_run(text)
            run.italic         = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Checkbox (- [ ] hoac - [x]) ────────────────────────────────────
        if re.match(r'^-\s+\[[ xX]\]', line):
            text = re.sub(r'^-\s+\[[ xX]\]\s*', '', line)
            p    = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, '[  ] ' + text)
            i += 1
            continue

        # ── Bullet list (ho tro thu lung 2 cap: 0, 2, 4 spaces) ─────────────
        m_bullet = re.match(r'^( *)[-\*]\s+(.*)', line)
        if m_bullet:
            indent_lvl = len(m_bullet.group(1)) // 2   # 0 = cap 1, 1 = cap 2
            text = m_bullet.group(2)
            p    = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(indent_lvl * 1.0)
            add_formatted_run(p, text)
            i += 1
            continue

        # ── Numbered list ──────────────────────────────────────────────────
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p    = doc.add_paragraph(style='List Number')
            add_formatted_run(p, text)
            i += 1
            continue

        # ── Math inline $$...$$ ────────────────────────────────────────────
        if line.strip().startswith('$$'):
            p              = doc.add_paragraph()
            run            = p.add_run(line.strip())
            run.font.name  = 'Cambria Math'
            run.font.size  = Pt(12)
            p.alignment    = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ── Dong trong ────────────────────────────────────────────────────
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # ── Doan van thuong ────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_formatted_run(p, line.strip())
        i += 1

    return doc


# === RUN ====================================================================
print("Bat dau chuyen doi MD -> DOCX voi Mermaid rendering...")
print(f"Input : {MD_FILE}")
print(f"Output: {DOCX_FILE}")
print()

doc = Document()
set_page_setup(doc)

# Xoa paragraph mac dinh rong
for para in doc.paragraphs:
    para._element.getparent().remove(para._element)

parse_and_write(doc, MD_FILE)
doc.save(DOCX_FILE)

print()
print(f"[OK] Hoan thanh! File: {DOCX_FILE}")
