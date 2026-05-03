import os
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from src.core.chapter_settings import ChapterSettings
from src.core.media_downloader import MediaDownloader
from src.core.docx_helpers import DocxHelpers, COLOR_H1, COLOR_H2, COLOR_H3, COLOR_H4
from src.core.markdown_utils import MarkdownUtils

class DocxBuilder:
    def __init__(self, workspace_dir: Path):
        self.workspace_dir = workspace_dir
        self.config = self._get_config()
        self.doc = self._init_document()
        self.diagram_idx = 0
        self.math_idx = 0

    def _get_config(self):
        from src.core.config import TemplateConfig
        config_path = self.workspace_dir / 'config.yaml'
        if config_path.exists():
            return TemplateConfig.load(config_path)
        return None

    def _init_document(self):
        """Loads template.docx if it exists, otherwise creates a blank Document."""
        page_settings = DocxHelpers.get_page_settings(self.config)
        if self.config and self.config.docx_template:
            template_path = self.workspace_dir / self.config.docx_template
            if template_path.exists():
                doc = Document(str(template_path))
                DocxHelpers.apply_page_settings(doc, page_settings)
                return doc
        doc = Document()
        DocxHelpers.apply_page_settings(doc, page_settings)
        return doc

    def build_from_markdown(self, md_path: str, img_cache_dir: Path):
        """
        Applies markdown parsing and writes to self.doc.
        Integrated from make.py's parse_and_write.
        """
        with open(md_path, encoding='utf-8') as f:
            lines = f.readlines()

        paragraph_settings = DocxHelpers.get_paragraph_settings(self.config)
        style = self.doc.styles['Normal']
        style.font.name = paragraph_settings.get('font_name', 'Times New Roman')
        style.font.size = Pt(float(paragraph_settings.get('font_size', 13)))
        style.paragraph_format.alignment = DocxHelpers.parse_alignment(paragraph_settings.get('alignment', 'justify'))
        style.paragraph_format.left_indent = Cm(float(paragraph_settings.get('left_indent_cm', 0.0)))
        style.paragraph_format.right_indent = Cm(float(paragraph_settings.get('right_indent_cm', 0.0)))
        style.paragraph_format.space_before = Pt(float(paragraph_settings.get('space_before_pt', 0.0)))
        style.paragraph_format.space_after = Pt(float(paragraph_settings.get('space_after_pt', 6.0)))
        special_indent = str(paragraph_settings.get('special_indent', 'first_line')).lower()
        special_indent_by_cm = float(paragraph_settings.get('special_indent_by_cm', 1.27))
        style.paragraph_format.first_line_indent = None
        if special_indent == 'first_line':
            style.paragraph_format.first_line_indent = Cm(special_indent_by_cm)
        elif special_indent == 'hanging':
            style.paragraph_format.first_line_indent = Cm(-special_indent_by_cm)

        line_spacing_mode = str(paragraph_settings.get('line_spacing_mode', 'multiple')).lower()
        line_spacing_value = float(paragraph_settings.get('line_spacing_value', 1.5))
        if line_spacing_mode == 'single':
            style.paragraph_format.line_spacing = 1.0
        elif line_spacing_mode == 'double':
            style.paragraph_format.line_spacing = 2.0
        elif line_spacing_mode == 'exactly':
            style.paragraph_format.line_spacing = Pt(line_spacing_value)
        else:
            style.paragraph_format.line_spacing = line_spacing_value

        # Tiền xử lý: rút gọn ≥2 dòng trống liên tiếp thành tối đa 1
        compressed, blank_run = [], 0
        for _ln in lines:
            if not _ln.strip():
                blank_run += 1
                if blank_run == 1:
                    compressed.append(_ln)
            else:
                blank_run = 0
                compressed.append(_ln)
        lines = compressed

        i, in_code, mermaid_block, mermaid_buf = 0, False, False, []
        current_source_filename = None

        while i < len(lines):
            line = lines[i].rstrip('\n')

            source_match = re.match(r'^\s*<!--\s*FILE:\s+(.+?)\s*-->\s*$', line)
            if source_match:
                current_source_filename = source_match.group(1).strip()
                i += 1
                continue

            # Code/Mermaid block
            if line.strip().startswith('```'):
                if not in_code:
                    in_code = True
                    lang = line.strip()[3:].strip().lower()
                    mermaid_block = lang in {'plantuml', 'puml'}
                    mermaid_buf = []
                else:
                    if mermaid_block and mermaid_buf:
                        self.diagram_idx += 1
                        img_path = MediaDownloader.render_plantuml('\n'.join(mermaid_buf), self.diagram_idx, str(img_cache_dir))
                        if img_path:
                            p = self.doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.space_before = Pt(6)
                            p.paragraph_format.space_after = Pt(6)
                            run = p.add_run()
                            max_width, max_height = DocxHelpers.get_content_frame_size(self.doc, height_reserve=Cm(3))
                            DocxHelpers.add_picture_fit(run, img_path, self.doc, max_width=max_width, max_height=max_height)
                            cap = self.doc.add_paragraph(f'Bieu do {self.diagram_idx}')
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.runs[0].font.size = Pt(10)
                            cap.runs[0].italic = True
                            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        else:
                            p = self.doc.add_paragraph()
                            r = p.add_run(f'[Bieu do PlantUML {self.diagram_idx} - khong the render]')
                            r.italic = True
                            r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                    in_code = mermaid_block = False
                    mermaid_buf = []
                i += 1
                continue

            if in_code:
                if mermaid_block:
                    mermaid_buf.append(line)
                else:
                    p = self.doc.add_paragraph()
                    run = p.add_run(line if line else ' ')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    DocxHelpers.set_para_shading(p, 'F4F4F4')
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                i += 1
                continue

            # Bảng
            if line.startswith('|'):
                table_rows = []
                while i < len(lines) and lines[i].startswith('|'):
                    cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                    table_rows.append(cells)
                    i += 1
                data_rows = [r for r in table_rows if not all(re.match(r'^[-: ]+$', c) for c in r)]
                if not data_rows:
                    continue
                max_cols = max(len(r) for r in data_rows)
                data_rows = [r + [''] * (max_cols - len(r)) for r in data_rows]
                tbl = self.doc.add_table(rows=len(data_rows), cols=max_cols)
                tbl.style = 'Table Grid'
                for ri, row in enumerate(data_rows):
                    for ci, cell_text in enumerate(row):
                        cell = tbl.cell(ri, ci)
                        cell.paragraphs[0].clear()
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run(MarkdownUtils.strip_md_markup(cell_text))
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        if ri == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            DocxHelpers.set_cell_bg(cell, '1F619E')
                        elif ri % 2 == 0:
                            DocxHelpers.set_cell_bg(cell, 'EBF4FB')
                _sp = self.doc.add_paragraph()
                _sp.paragraph_format.space_before = Pt(0)
                _sp.paragraph_format.space_after = Pt(6)
                _sp.add_run('').font.size = Pt(4)
                continue

            # Horizontal rule
            if re.match(r'^---+\s*$', line):
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                i += 1
                continue

            # Centered text
            m_center = re.match(r'^\s*<center>(.*?)</center>\s*$', line)
            if m_center:
                text = m_center.group(1).strip()
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                MarkdownUtils.add_formatted_run(p, text)
                i += 1
                continue

            # Markdown image
            m_image = re.match(r'^\s*!\[.*?\]\(([^)]+)\)\s*$', line)
            if m_image:
                DocxHelpers.add_markdown_image(self.doc, self.workspace_dir, md_path, m_image.group(1).strip())
                i += 1
                continue

            if line.strip() == '[[PAGEBREAK]]':
                DocxHelpers.add_page_break(self.doc)
                i += 1
                continue

            if line.strip() == '[[TOC]]':
                DocxHelpers.add_table_of_contents(self.doc)
                i += 1
                continue

            # Heading
            m = re.match(r'^(#{1,6})\s+(.*)', line)
            if m:
                level = len(m.group(1))
                text = MarkdownUtils.strip_md_markup(m.group(2))
                p = self.doc.add_paragraph(style=f'Heading {min(level, 4)}')
                heading_indent = MarkdownUtils.get_heading_indent(text)
                if heading_indent is not None:
                    p.paragraph_format.left_indent = heading_indent
                _space_before = [Pt(12), Pt(10), Pt(8), Pt(6), Pt(6), Pt(4)]
                p.paragraph_format.space_before = _space_before[level - 1]
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = [COLOR_H1, COLOR_H2, COLOR_H3, COLOR_H4, COLOR_H4, COLOR_H4][level - 1]
                run.font.size = Pt([16, 14, 13, 12, 12, 11][level - 1])
                run.bold = level <= 3
                i += 1
                continue

            # Blockquote
            if line.startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].startswith('>'):
                    quote_lines.append(lines[i].lstrip('> ').strip())
                    i += 1

                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                MarkdownUtils.add_formatted_run(p, ' '.join(quote_lines))

                for run in p.runs:
                    if run.font.name != 'Courier New':
                        run.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                continue

            # Checkbox
            if re.match(r'^-\s+\[[ xX]\]', line):
                text = re.sub(r'^-\s+\[[ xX]\]\s*', '', line)
                p = self.doc.add_paragraph()
                base_indent_cm = float(paragraph_settings.get('left_indent_cm', 0.0))
                if str(paragraph_settings.get('special_indent', 'first_line')).lower() == 'first_line':
                    base_indent_cm += float(paragraph_settings.get('special_indent_by_cm', 1.27))
                p.paragraph_format.left_indent = Cm(base_indent_cm)
                p.paragraph_format.first_line_indent = Cm(0.0)
                MarkdownUtils.add_formatted_run(p, '[  ] ' + text)
                i += 1
                continue

            # Bullet list
            m_bullet = re.match(r'^( *)[-\*\+]\s+(.*)', line)
            if m_bullet:
                indent_lvl = len(m_bullet.group(1)) // 2
                marker = ChapterSettings.get_list_marker(self.config, current_source_filename or '', indent_lvl + 1)
                p = self.doc.add_paragraph()
                base_indent_cm = float(paragraph_settings.get('left_indent_cm', 0.0))
                if str(paragraph_settings.get('special_indent', 'first_line')).lower() == 'first_line':
                    base_indent_cm += float(paragraph_settings.get('special_indent_by_cm', 1.27))
                p.paragraph_format.left_indent = Cm(base_indent_cm + (indent_lvl * 1.0))
                p.paragraph_format.first_line_indent = Cm(0.0)
                MarkdownUtils.add_formatted_run(p, f'{marker} {m_bullet.group(2)}')
                i += 1
                continue

            # Numbered list
            if re.match(r'^\d+\.\s+', line):
                text = re.sub(r'^\d+\.\s+', '', line)
                p = self.doc.add_paragraph(style='List Number')
                MarkdownUtils.add_formatted_run(p, text)
                i += 1
                continue

            # Công thức toán (khối $$ ... $$)
            if line.strip() == '$$':
                math_buf = []
                i += 1
                while i < len(lines) and lines[i].rstrip('\n').strip() != '$$':
                    math_buf.append(lines[i].rstrip('\n'))
                    i += 1
                i += 1
                latex_code = ' '.join(b.strip() for b in math_buf if b.strip())
                if latex_code:
                    self.math_idx += 1
                    img_path = MediaDownloader.render_latex(latex_code, self.math_idx, str(img_cache_dir))
                    if img_path:
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        max_width, max_height = DocxHelpers.get_content_frame_size(self.doc, height_reserve=Cm(4))
                        DocxHelpers.add_picture_fit(p.add_run(), img_path, self.doc, max_width=max_width, max_height=max_height)
                    else:
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(latex_code)
                        run.font.name = 'Cambria Math'
                        run.font.size = Pt(11)
                continue

            # Công thức toán inline: $$ ... $$ trên một dòng
            if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
                latex_code = line.strip()[2:-2].strip()
                if latex_code:
                    self.math_idx += 1
                    img_path = MediaDownloader.render_latex(latex_code, self.math_idx, str(img_cache_dir))
                    if img_path:
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        max_width, max_height = DocxHelpers.get_content_frame_size(self.doc, height_reserve=Cm(4))
                        DocxHelpers.add_picture_fit(p.add_run(), img_path, self.doc, max_width=max_width, max_height=max_height)
                    else:
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run(latex_code).font.name = 'Cambria Math'
                i += 1
                continue

            # Dòng trống → bỏ qua hoàn toàn, không tạo paragraph
            if not line.strip():
                i += 1
                continue

            # Đoạn văn thường — chuẩn NĐ30: justify, lùi đầu dòng, dãn 1.5
            p = self.doc.add_paragraph()
            DocxHelpers.apply_paragraph_format(p, paragraph_settings)
            MarkdownUtils.add_formatted_run(p, MarkdownUtils.normalize_punctuation(line.strip()))
            i += 1

    def save(self, output_path: Path):
        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            self.doc.save(str(output_path))
        except PermissionError as exc:
            raise RuntimeError(
                f"Khong the ghi file {output_path.name}. Hay dong file Word cu truoc khi build lai."
            ) from exc
