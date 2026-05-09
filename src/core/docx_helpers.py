import re
import unicodedata
from pathlib import Path

from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

from src.core.markdown_image import MarkdownImage
from src.core.media_downloader import MediaDownloader

# ── MÀUSẮC ───────────────────────────────────────────────────────────────────
COLOR_H1 = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_H2 = RGBColor(0x1F, 0x61, 0x9E)
COLOR_H3 = RGBColor(0x2E, 0x86, 0xAB)
COLOR_H4 = RGBColor(0x44, 0x9D, 0xD1)


class DocxHelpers:
    EXAM_COVER_LEGACY_TEXT_MAP = {
        'HOC VIEN CONG NGHE BUU CHINH VIEN THONG': 'HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG',
        'KHOA CONG NGHE THONG TIN 1': 'KHOA CÔNG NGHỆ THÔNG TIN 1',
        'BAI KIEM TRA GIUA KY': 'BÀI KIỂM TRA GIỮA KỲ',
        'MON: CO SO DU LIEU PHAN TAN': 'MÔN: CƠ SỞ DỮ LIỆU PHÂN TÁN',
        'MON: ...': 'MÔN: ...',
        'Giang vien': 'Giảng viên',
        'Lop': 'Lớp',
        'Ho ten': 'Họ tên',
        'Ha Noi, Thang .../....': 'Hà Nội, Tháng .../....',
        'Ha Noi, Thang 05/2026': 'Hà Nội, Tháng 05/2026',
    }

    @staticmethod
    def get_exam_cover_settings(config):
        settings = (config.settings if config and config.settings else {}).copy()
        cover = (settings.get('cover_page', {}) or {}).copy()
        defaults = {
            'institution_line_1': 'HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG',
            'institution_line_2': 'KHOA CÔNG NGHỆ THÔNG TIN 1',
            'divider': '______________________________',
            'logo_path': 'assets/ptit-logo.png',
            'title': 'BÀI KIỂM TRA GIỮA KỲ',
            'subject': 'MÔN: ...',
            'lecturer_label': 'Giảng viên',
            'lecturer_value': '',
            'class_label': 'Lớp',
            'class_value': '',
            'student_label': 'Họ tên',
            'student_value': '',
            'student_id_label': 'MSSV',
            'student_id_value': '',
            'date_text': '',
        }
        defaults.update(cover)
        for key, value in list(defaults.items()):
            if isinstance(value, str):
                defaults[key] = DocxHelpers.EXAM_COVER_LEGACY_TEXT_MAP.get(value, value)
        return defaults

    @staticmethod
    def parse_exam_cover_content(content, defaults):
        cover = dict(defaults)
        mapping = {
            'TIEUDE': 'title',
            'MON': 'subject',
            'GIANGVIEN': 'lecturer_value',
            'LOP': 'class_value',
            'HOTEN': 'student_value',
            'MSSV': 'student_id_value',
            'THOIGIAN': 'date_text',
        }

        def normalize_key(value):
            normalized = unicodedata.normalize('NFD', value)
            stripped = ''.join(char for char in normalized if unicodedata.category(char) != 'Mn')
            return re.sub(r'[^A-Z]', '', stripped.upper())

        for raw_line in content.splitlines():
            line = raw_line.strip()
            if not line or ':' not in line:
                continue
            left, value = line.split(':', 1)
            key = normalize_key(left)
            value = value.strip().strip('*').strip()
            target = mapping.get(key)
            if target and value:
                cover[target] = value
        return cover

    @staticmethod
    def get_markdown_table_settings(config):
        settings = (config.settings if config and config.settings else {}).copy()
        table_settings = (settings.get('markdown_table', {}) or {}).copy()
        defaults = {
            'font_name': 'Times New Roman',
            'font_size': 11,
            'header_fill': 'D9EAF7',
            'header_alignment': 'center',
            'body_alignment': 'left',
            'vertical_alignment': 'center',
            'space_before_pt': 2.0,
            'space_after_pt': 2.0,
            'line_spacing': 1.0,
            'cell_padding_top_dxa': 60,
            'cell_padding_bottom_dxa': 60,
            'cell_padding_left_dxa': 108,
            'cell_padding_right_dxa': 108,
        }
        defaults.update(table_settings)
        return defaults

    @staticmethod
    def get_page_settings(config):
        settings = (config.settings if config and config.settings else {}).copy()
        defaults = {
            'page_width_cm': 21.0,
            'page_height_cm': 29.7,
            'margin_top_cm': 2.5,
            'margin_bottom_cm': 2.5,
            'margin_left_cm': 3.0,
            'margin_right_cm': 2.0,
            'gutter_cm': 0.0,
            'orientation': 'portrait',
        }
        defaults.update(settings)
        return defaults

    @staticmethod
    def get_paragraph_settings(config):
        settings = (config.settings if config and config.settings else {}).copy()
        defaults = {
            'font_name': 'Times New Roman',
            'font_size': 14,
            'alignment': 'justify',
            'left_indent_cm': 0.0,
            'right_indent_cm': 0.0,
            'special_indent': 'first_line',
            'special_indent_by_cm': 1.27,
            'space_before_pt': 0.0,
            'space_after_pt': 6.0,
            'line_spacing_mode': 'multiple',
            'line_spacing_value': 1.5,
        }
        defaults.update(settings)
        return defaults

    @staticmethod
    def parse_alignment(value):
        mapping = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return mapping.get(str(value).lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)

    @staticmethod
    def apply_paragraph_format(paragraph, settings, use_spacing_after=True):
        paragraph.alignment = DocxHelpers.parse_alignment(settings.get('alignment', 'justify'))
        paragraph.paragraph_format.left_indent = Cm(float(settings.get('left_indent_cm', 0.0)))
        paragraph.paragraph_format.right_indent = Cm(float(settings.get('right_indent_cm', 0.0)))

        special_indent = str(settings.get('special_indent', 'first_line')).lower()
        special_indent_by_cm = float(settings.get('special_indent_by_cm', 1.27))
        paragraph.paragraph_format.first_line_indent = None
        if special_indent == 'first_line':
            paragraph.paragraph_format.first_line_indent = Cm(special_indent_by_cm)
        elif special_indent == 'hanging':
            paragraph.paragraph_format.first_line_indent = Cm(-special_indent_by_cm)

        paragraph.paragraph_format.space_before = Pt(float(settings.get('space_before_pt', 0.0)))
        paragraph.paragraph_format.space_after = Pt(
            float(settings.get('space_after_pt', 6.0 if use_spacing_after else 0.0))
        )

        line_spacing_mode = str(settings.get('line_spacing_mode', 'multiple')).lower()
        line_spacing_value = float(settings.get('line_spacing_value', 1.5))
        if line_spacing_mode == 'single':
            paragraph.paragraph_format.line_spacing = 1.0
        elif line_spacing_mode == 'double':
            paragraph.paragraph_format.line_spacing = 2.0
        elif line_spacing_mode == 'exactly':
            paragraph.paragraph_format.line_spacing = Pt(line_spacing_value)
        else:
            paragraph.paragraph_format.line_spacing = line_spacing_value

    @staticmethod
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    @staticmethod
    def set_para_shading(para, hex_color):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        pPr.append(shd)

    @staticmethod
    def set_page_setup(doc):
        DocxHelpers.apply_page_settings(doc, DocxHelpers.get_page_settings(None))

    @staticmethod
    def apply_page_settings(doc, settings):
        sec = doc.sections[0]
        width_cm = float(settings.get('page_width_cm', 21.0))
        height_cm = float(settings.get('page_height_cm', 29.7))
        orientation = str(settings.get('orientation', 'portrait')).lower()
        if orientation == 'landscape':
            sec.orientation = WD_ORIENTATION.LANDSCAPE
            sec.page_width = Cm(max(width_cm, height_cm))
            sec.page_height = Cm(min(width_cm, height_cm))
        else:
            sec.orientation = WD_ORIENTATION.PORTRAIT
            sec.page_width = Cm(min(width_cm, height_cm))
            sec.page_height = Cm(max(width_cm, height_cm))

        sec.left_margin = Cm(float(settings.get('margin_left_cm', 3.0)))
        sec.right_margin = Cm(float(settings.get('margin_right_cm', 2.0)))
        sec.top_margin = Cm(float(settings.get('margin_top_cm', 2.5)))
        sec.bottom_margin = Cm(float(settings.get('margin_bottom_cm', 2.5)))
        sec.gutter = Cm(float(settings.get('gutter_cm', 0.0)))

    @staticmethod
    def get_content_frame_size(doc, height_reserve=Cm(1.5)):
        sec = doc.sections[0]
        # python-docx trả về int thô (EMU) khi trừ 2 Length objects.
        # Cần bọc lại bằng Emu() để có .inches attribute cho tính toán scale.
        max_width = Emu(sec.page_width - sec.left_margin - sec.right_margin)
        max_height = Emu(sec.page_height - sec.top_margin - sec.bottom_margin - height_reserve)
        return max_width, max_height

    @staticmethod
    def add_picture_fit(run, image_path, doc, max_width=None, max_height=None):
        if max_width is None or max_height is None:
            max_width, max_height = DocxHelpers.get_content_frame_size(doc)

        img_w, img_h = MediaDownloader.get_image_dimensions(image_path)
        if not img_w or not img_h:
            run.add_picture(str(image_path), width=max_width)
            return

        width_inches = img_w / 96.0
        height_inches = img_h / 96.0

        width_scale = max_width.inches / width_inches if width_inches else 1.0
        height_scale = max_height.inches / height_inches if height_inches else 1.0
        scale = min(width_scale, height_scale)

        target_width = Inches(width_inches * scale)
        target_height = Inches(height_inches * scale)
        run.add_picture(str(image_path), width=target_width, height=target_height)

    @staticmethod
    def parse_image_width_fraction(width: str) -> float:
        value = str(width or '100%').strip()
        if value.endswith('%'):
            try:
                return max(0.2, min(1.0, float(value[:-1]) / 100.0))
            except ValueError:
                return 1.0
        return 1.0

    @staticmethod
    def parse_image_alignment(value: str):
        alignment = str(value or 'center').strip().lower()
        if alignment == 'left':
            return WD_ALIGN_PARAGRAPH.LEFT
        if alignment == 'right':
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def configure_media_paragraph(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=18):
        paragraph.alignment = alignment
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        # Inline images behave like oversized glyphs in Word. If the base style uses
        # exact line spacing, the image can visually overlap nearby paragraphs unless
        # the media paragraph explicitly resets line spacing to auto/single.
        paragraph.paragraph_format.line_spacing = 1.0

    @staticmethod
    def set_table_borders(table, *, visible):
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in('w:tblBorders')
        if borders is None:
            borders = OxmlElement('w:tblBorders')
            tbl_pr.append(borders)

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            element = borders.find(qn(f'w:{edge}'))
            if element is None:
                element = OxmlElement(f'w:{edge}')
                borders.append(element)
            if visible:
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), 'auto')
            else:
                element.set(qn('w:val'), 'nil')

    @staticmethod
    def set_table_row_repeat_header(row):
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn('w:tblHeader'))
        if tbl_header is None:
            tbl_header = OxmlElement('w:tblHeader')
            tr_pr.append(tbl_header)
        tbl_header.set(qn('w:val'), 'true')

    @staticmethod
    def set_cell_margins(cell, *, top_dxa, right_dxa, bottom_dxa, left_dxa):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn('w:tcMar'))
        if tc_mar is None:
            tc_mar = OxmlElement('w:tcMar')
            tc_pr.append(tc_mar)

        for edge, value in (
            ('top', top_dxa),
            ('right', right_dxa),
            ('bottom', bottom_dxa),
            ('left', left_dxa),
        ):
            element = tc_mar.find(qn(f'w:{edge}'))
            if element is None:
                element = OxmlElement(f'w:{edge}')
                tc_mar.append(element)
            element.set(qn('w:w'), str(int(value)))
            element.set(qn('w:type'), 'dxa')

    @staticmethod
    def format_markdown_table_cell(cell, text, *, is_header, settings):
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = DocxHelpers.parse_alignment(
            settings.get('header_alignment', 'center') if is_header else settings.get('body_alignment', 'left')
        )
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(float(settings.get('space_before_pt', 2.0)))
        paragraph.paragraph_format.space_after = Pt(float(settings.get('space_after_pt', 2.0)))
        paragraph.paragraph_format.line_spacing = float(settings.get('line_spacing', 1.0))

        vertical_alignment = str(settings.get('vertical_alignment', 'center')).lower()
        if vertical_alignment == 'center':
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif vertical_alignment == 'bottom':
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        else:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        DocxHelpers.set_cell_margins(
            cell,
            top_dxa=settings.get('cell_padding_top_dxa', 60),
            right_dxa=settings.get('cell_padding_right_dxa', 108),
            bottom_dxa=settings.get('cell_padding_bottom_dxa', 60),
            left_dxa=settings.get('cell_padding_left_dxa', 108),
        )

        if is_header:
            DocxHelpers.set_cell_bg(cell, settings.get('header_fill', 'D9EAF7'))

        run = paragraph.add_run(text)
        run.font.size = Pt(float(settings.get('font_size', 11)))
        run.font.name = settings.get('font_name', 'Times New Roman')
        run.bold = bool(is_header)
        return paragraph

    @staticmethod
    def resolve_cover_asset_path(workspace_dir, relative_path):
        if not relative_path:
            return None
        candidate = Path(relative_path)
        if candidate.is_absolute():
            return candidate if candidate.exists() else None
        resolved = (Path(workspace_dir) / candidate).resolve()
        return resolved if resolved.exists() else None

    @staticmethod
    def add_exam_cover(doc, workspace_dir, config):
        cover = DocxHelpers.get_exam_cover_settings(config)

        def add_centered(text, *, font_size=14, bold=False, italic=False, spacing_before=0, spacing_after=0):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(spacing_before)
            paragraph.paragraph_format.space_after = Pt(spacing_after)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run(text)
            run.font.name = cover.get('font_name', 'Times New Roman')
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            return paragraph

        add_centered(cover['institution_line_1'], font_size=14, bold=True, spacing_after=6)
        add_centered(cover['institution_line_2'], font_size=14, bold=True, spacing_after=24)
        add_centered(cover['divider'], font_size=12, spacing_after=18)

        logo_path = DocxHelpers.resolve_cover_asset_path(workspace_dir, cover.get('logo_path'))
        if logo_path:
            logo_paragraph = doc.add_paragraph()
            DocxHelpers.configure_media_paragraph(logo_paragraph, space_before=0, space_after=26)
            logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(1.8))

        add_centered(cover['title'], font_size=18, bold=True, spacing_after=28)
        add_centered(cover['subject'], font_size=16, bold=True, spacing_after=36)

        table = doc.add_table(rows=4, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        DocxHelpers.set_table_borders(table, visible=False)

        widths = (Cm(4.2), Cm(0.8), Cm(7.8))
        rows = [
            (cover['lecturer_label'], ':', cover['lecturer_value']),
            (cover['class_label'], ':', cover['class_value']),
            (cover['student_label'], ':', cover['student_value']),
            (cover['student_id_label'], ':', cover['student_id_value']),
        ]
        for row_index, values in enumerate(rows):
            for col_index, value in enumerate(values):
                cell = table.rows[row_index].cells[col_index]
                cell.width = widths[col_index]
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 1 else WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(value)
                run.font.name = cover.get('font_name', 'Times New Roman')
                run.font.size = Pt(14)
                if col_index != 1:
                    run.bold = True

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(160)
        spacer.paragraph_format.line_spacing = 1.0

        add_centered(cover['date_text'], font_size=13, italic=True, spacing_after=0)
        DocxHelpers.add_page_break(doc)

    @staticmethod
    def set_picture_wrap_top_bottom(run, pic_id=1):
        """Chuyen hinh anh inline (wp:inline) thanh floating Top and Bottom (wp:anchor).
        Hinh duoc can giua theo chieu ngang, nam dung vi tri dat vao van ban.
        """
        r_elem = run._r
        drawing = r_elem.find(qn('w:drawing'))
        if drawing is None:
            return
        inline = drawing.find(qn('wp:inline'))
        if inline is None:
            return  # da la anchor hoac khong co

        # Lay kich thuoc tu wp:extent
        extent = inline.find(qn('wp:extent'))
        cx = extent.get('cx', '5000000') if extent is not None else '5000000'
        cy = extent.get('cy', '3000000') if extent is not None else '3000000'

        # Lay doc properties
        doc_pr = inline.find(qn('wp:docPr'))
        pid = doc_pr.get('id', str(pic_id)) if doc_pr is not None else str(pic_id)
        pname = doc_pr.get('name', f'Image {pic_id}') if doc_pr is not None else f'Image {pic_id}'
        pdescr = doc_pr.get('descr', '') if doc_pr is not None else ''

        # Lay cNvGraphicFramePr va graphic — dung namespace day du
        cnv = inline.find(qn('wp:cNvGraphicFramePr'))
        graphic = inline.find(qn('a:graphic'))

        # --- Xay dung wp:anchor ---
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', '114300')  # 0.9 cm tren
        anchor.set('distB', '114300')  # 0.9 cm duoi
        anchor.set('distL', '0')
        anchor.set('distR', '0')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251659264')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '0')

        sp = OxmlElement('wp:simplePos')
        sp.set('x', '0')
        sp.set('y', '0')
        anchor.append(sp)

        ph = OxmlElement('wp:positionH')
        ph.set('relativeFrom', 'column')
        ah = OxmlElement('wp:align')
        ah.text = 'center'
        ph.append(ah)
        anchor.append(ph)

        pv = OxmlElement('wp:positionV')
        pv.set('relativeFrom', 'paragraph')
        po = OxmlElement('wp:posOffset')
        po.text = '0'
        pv.append(po)
        anchor.append(pv)

        ext2 = OxmlElement('wp:extent')
        ext2.set('cx', cx)
        ext2.set('cy', cy)
        anchor.append(ext2)

        ee = OxmlElement('wp:effectExtent')
        for attr in ('l', 't', 'r', 'b'):
            ee.set(attr, '0')
        anchor.append(ee)

        anchor.append(OxmlElement('wp:wrapTopAndBottom'))

        dp = OxmlElement('wp:docPr')
        dp.set('id', pid)
        dp.set('name', pname)
        if pdescr:
            dp.set('descr', pdescr)
        anchor.append(dp)

        if cnv is not None:
            anchor.append(cnv)
        else:
            anchor.append(OxmlElement('wp:cNvGraphicFramePr'))

        # QUAN TRONG: graphic phai co mat, neu khong Word tu choi mo file
        if graphic is not None:
            anchor.append(graphic)
        else:
            # Fallback: abort conversion, giu nguyen inline
            return

        drawing.replace(inline, anchor)

    @staticmethod
    def resolve_media_path(base_dir, md_path, asset_path):
        asset = Path(asset_path)
        if asset.is_absolute():
            return asset
        primary = (Path(md_path).resolve().parent / asset).resolve()
        if primary.exists():
            return primary
        fallback_text = asset_path.replace('../', '').replace('..\\', '').replace('./', '').replace('.\\', '')
        return (Path(base_dir) / Path(fallback_text)).resolve()

    @staticmethod
    def add_markdown_image(doc, base_dir, md_path, image):
        if isinstance(image, MarkdownImage):
            image_ref = image.path
            caption = image.caption
            width_fraction = DocxHelpers.parse_image_width_fraction(image.width)
            alignment = DocxHelpers.parse_image_alignment(image.align)
        else:
            image_ref = str(image)
            caption = ''
            width_fraction = 1.0
            alignment = WD_ALIGN_PARAGRAPH.CENTER

        image_path = DocxHelpers.resolve_media_path(base_dir, md_path, image_ref)
        if not image_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'[Khong tim thay anh: {image_ref}]')
            r.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            return

        p = doc.add_paragraph()
        DocxHelpers.configure_media_paragraph(p, alignment=alignment)
        run = p.add_run()
        max_width, max_height = DocxHelpers.get_content_frame_size(doc, height_reserve=Cm(2))
        max_width = Emu(int(max_width * width_fraction))
        DocxHelpers.add_picture_fit(run, image_path, doc, max_width=max_width, max_height=max_height)
        if caption:
            caption_paragraph = doc.add_paragraph()
            caption_paragraph.alignment = alignment
            caption_paragraph.paragraph_format.left_indent = Cm(0)
            caption_paragraph.paragraph_format.right_indent = Cm(0)
            caption_paragraph.paragraph_format.first_line_indent = Cm(0)
            caption_paragraph.paragraph_format.space_before = Pt(0)
            caption_paragraph.paragraph_format.space_after = Pt(8)
            caption_paragraph.paragraph_format.line_spacing = 1.0
            run = caption_paragraph.add_run(caption)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True

    @staticmethod
    def add_page_break(doc):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    @staticmethod
    def add_table_of_contents(doc):
        p = doc.add_paragraph()
        begin_run = OxmlElement('w:r')
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        begin_run.append(fld_begin)

        instr_run = OxmlElement('w:r')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-4" \\h \\z \\u'
        instr_run.append(instr)

        separate_run = OxmlElement('w:r')
        fld_separate = OxmlElement('w:fldChar')
        fld_separate.set(qn('w:fldCharType'), 'separate')
        separate_run.append(fld_separate)

        hint_run = OxmlElement('w:r')
        hint_text = OxmlElement('w:t')
        hint_text.text = 'Cap nhat Muc luc trong Word bang Update Field.'
        hint_run.append(hint_text)

        end_run = OxmlElement('w:r')
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        end_run.append(fld_end)

        p._p.append(begin_run)
        p._p.append(instr_run)
        p._p.append(separate_run)
        p._p.append(hint_run)
        p._p.append(end_run)
