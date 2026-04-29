"""
make.py — Tool duy nhat: Ghep chapters -> MD -> convert DOCX
Usage: uv run python make.py
"""

import argparse
import html as html_lib
import os
import re
import shutil
import struct
import sys
import time
from html.parser import HTMLParser
from pathlib import Path
from urllib.parse import quote

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

import split_chapters
from convert_docx_to_md import convert_docx_to_markdown

sys.stdout.reconfigure(encoding='utf-8')

# ── CONFIG ───────────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent
CH_DIR = str(BASE / 'chapters')
MD_OUT = str(BASE / 'Bao_Cao_Tieu_Luan_NMCNPM.md')
DOCX_OUT = str(BASE / 'Bao_Cao_Tieu_Luan_NMCNPM.docx')
IMG_CACHE = str(BASE / 'diagram_cache')
os.makedirs(IMG_CACHE, exist_ok=True)

# Danh sách chapter chính thức dùng để build báo cáo.
# Các file Ch*.md khác trong thư mục chapters được giữ lại để tham chiếu,
# nhưng không được ghép vào báo cáo để tránh trùng/lặp chương.
BUILD_CHAPTER_FILES = [
    'F00_header.md',
    'F01_toc.md',
    'Ch01_LỜI_MỞ_ĐẦU.md',
    'Ch02_MỞ_ĐẦU_KẾ_HOẠCH_QUẢN_LÝ_DỰ_ÁN_PHẦN_MỀM_.md',
    'Ch03_CHƯƠNG_1_TỔNG_QUAN_HỆ_THỐNG_VÀ_PHÂN_CÔN.md',
    'Ch04_CHƯƠNG_2_THIẾT_KẾ_KIẾN_TRÚC_VÀ_CƠ_SỞ_DỮ.md',
    'Ch05_CHƯƠNG_3_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch06_CHƯƠNG_4_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch07_CHƯƠNG_5_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch08_CHƯƠNG_6_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch09_CHƯƠNG_7_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch10_CHƯƠNG_8_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch11_CHƯƠNG_9_NGHIÊN_CỨU_CHUYÊN_SÂU_—_USE_CA.md',
    'Ch12_KẾT_LUẬN.md',
    'Ch13_TÀI_LIỆU_THAM_KHẢO.md',
]

# ── MÀUSẮC ───────────────────────────────────────────────────────────────────
COLOR_H1 = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_H2 = RGBColor(0x1F, 0x61, 0x9E)
COLOR_H3 = RGBColor(0x2E, 0x86, 0xAB)
COLOR_H4 = RGBColor(0x44, 0x9D, 0xD1)


# ════════════════════════════════════════════════════════════════════════════
# BƯỚC 1: GỘP CHAPTERS → MD
# ════════════════════════════════════════════════════════════════════════════
def chapter_sort_key(file_path):
    name = os.path.basename(file_path)
    lower_name = name.lower()

    if lower_name == 'f00_header.md':
        return (0, 0, '')
    if lower_name == 'f01_toc.md':
        return (0, 1, '')

    match = re.match(r'^Ch(\d+)(.*)\.md$', name, re.IGNORECASE)
    if not match:
        return (99, 99, lower_name)

    return (1, int(match.group(1)), match.group(2).lower())


def collect_chapter_files(chapter_dir):
    chapter_dir_path = Path(chapter_dir)
    chapter_files = []
    missing_files = []

    for filename in BUILD_CHAPTER_FILES:
        file_path = chapter_dir_path / filename
        if file_path.exists():
            chapter_files.append(str(file_path))
        else:
            missing_files.append(filename)

    if missing_files:
        missing_text = ', '.join(missing_files)
        raise FileNotFoundError(f'Thieu chapter trong danh sach build: {missing_text}')

    return chapter_files


def clear_image_cache(img_cache):
    cache_path = Path(img_cache)
    if not cache_path.exists():
        cache_path.mkdir(parents=True, exist_ok=True)
        return

    for entry in cache_path.iterdir():
        if entry.is_file():
            entry.unlink()
        elif entry.is_dir():
            shutil.rmtree(entry)


def ensure_output_docx_closed(docx_out):
    output_path = Path(docx_out)
    if not output_path.exists():
        return

    probe_path = output_path.with_name(output_path.stem + '.__lockcheck__' + output_path.suffix)
    try:
        os.replace(output_path, probe_path)
        os.replace(probe_path, output_path)
    except PermissionError as exc:
        if probe_path.exists():
            try:
                os.replace(probe_path, output_path)
            except OSError:
                pass
        raise RuntimeError(
            f'File output dang mo: {output_path.name}. Hay dong file Word cu truoc khi build lai.'
        ) from exc


def assemble_markdown(chapter_dir, output_path):
    chapter_files = collect_chapter_files(chapter_dir)
    parts = []

    for file_path in chapter_files:
        with open(file_path, encoding='utf-8') as handle:
            content = handle.read().strip()
        parts.append(content)
        print(f'  [OK] {os.path.basename(file_path)}  ({len(content.splitlines())} dong)')

    final = '\n\n'.join(parts)
    with open(output_path, 'w', encoding='utf-8') as handle:
        handle.write(final)

    return final, chapter_files


def step_assemble():
    print('=' * 55)
    print('BƯỚC 1: Ghép chapters → MD')
    print('=' * 55)

    final, chapter_files = assemble_markdown(CH_DIR, MD_OUT)

    kb = len(final.encode('utf-8')) // 1024
    print(f'\n  => {MD_OUT}')
    print(f'     {len(final.splitlines())} dong | {kb} KB | {len(chapter_files)} chapters\n')


# ════════════════════════════════════════════════════════════════════════════
# BƯỚC 2: CONVERT MD → DOCX
# ════════════════════════════════════════════════════════════════════════════


# ── Diagram render ───────────────────────────────────────────────────────────
def get_png_dimensions(path):
    with open(path, 'rb') as f:
        f.read(16)
        w = struct.unpack('>I', f.read(4))[0]
        h = struct.unpack('>I', f.read(4))[0]
    return w, h


def get_jpeg_dimensions(path):
    with open(path, 'rb') as f:
        if f.read(2) != b'\xff\xd8':
            raise ValueError('Not a JPEG file')

        while True:
            marker_prefix = f.read(1)
            if not marker_prefix:
                break
            if marker_prefix != b'\xff':
                continue

            marker = f.read(1)
            while marker == b'\xff':
                marker = f.read(1)

            if not marker or marker in {
                b'\xd8',
                b'\xd9',
                b'\x01',
                b'\xd0',
                b'\xd1',
                b'\xd2',
                b'\xd3',
                b'\xd4',
                b'\xd5',
                b'\xd6',
                b'\xd7',
            }:
                continue

            segment_length = struct.unpack('>H', f.read(2))[0]
            if marker in {
                b'\xc0',
                b'\xc1',
                b'\xc2',
                b'\xc3',
                b'\xc5',
                b'\xc6',
                b'\xc7',
                b'\xc9',
                b'\xca',
                b'\xcb',
                b'\xcd',
                b'\xce',
                b'\xcf',
            }:
                f.read(1)  # precision
                height = struct.unpack('>H', f.read(2))[0]
                width = struct.unpack('>H', f.read(2))[0]
                return width, height

            f.seek(segment_length - 2, os.SEEK_CUR)

    raise ValueError('Could not determine JPEG dimensions')


def get_image_dimensions(path):
    suffix = Path(path).suffix.lower()
    if suffix == '.png':
        return get_png_dimensions(path)
    if suffix in {'.jpg', '.jpeg'}:
        return get_jpeg_dimensions(path)
    return None, None


def plantuml_hex_encode(text):
    return '~h' + text.encode('utf-8').hex()


def plantuml_png_url(code):
    normalized = code.strip()
    if '@startuml' not in normalized:
        normalized = f'@startuml\n{normalized}\n@enduml'
    return f'https://www.plantuml.com/plantuml/png/{plantuml_hex_encode(normalized)}'


def render_plantuml(code, idx, img_cache):
    cache_file = os.path.join(img_cache, f'diagram_{idx:03d}.png')
    if os.path.exists(cache_file):
        print(f'  [cache] diagram_{idx:03d}.png')
        return cache_file
    try:
        print(f'  [render] Dang render diagram {idx}...')
        url = plantuml_png_url(code)
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200 and resp.headers.get('content-type', '').startswith('image'):
            with open(cache_file, 'wb') as f:
                f.write(resp.content)
            print(f'  [OK] Luu vao {cache_file}')
            time.sleep(0.5)
            return cache_file
        print(f'  [WARN] API tra ve {resp.status_code}')
        return None
    except Exception as e:
        print(f'  [ERROR] {e}')
        return None


def render_latex(latex_code, idx, img_cache):
    """Render công thức LaTeX → PNG qua CodeCogs API, cache lại kết quả."""
    cache_file = os.path.join(img_cache, f'math_{idx:03d}.png')
    if os.path.exists(cache_file):
        print(f'  [cache] math_{idx:03d}.png')
        return cache_file
    try:
        print(f'  [render] Dang render công thức toán {idx}...')
        encoded = quote(latex_code, safe='')
        url = f'https://latex.codecogs.com/png.image?\\dpi{{150}}{encoded}'
        resp = requests.get(url, timeout=15)
        if resp.status_code == 200 and resp.headers.get('content-type', '').startswith('image'):
            with open(cache_file, 'wb') as f:
                f.write(resp.content)
            print(f'  [OK] math_{idx:03d}.png')
            return cache_file
        print(f'  [WARN] CodeCogs trả về {resp.status_code}')
        return None
    except Exception as e:
        print(f'  [ERROR] Render LaTeX: {e}')
        return None


# ── DOCX helpers ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_page_setup(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)


def get_content_frame_size(doc, height_reserve=Cm(1.5)):
    sec = doc.sections[0]
    max_width = Emu(sec.page_width - sec.left_margin - sec.right_margin)
    max_height = Emu(sec.page_height - sec.top_margin - sec.bottom_margin - height_reserve)
    return max_width, max_height


def add_picture_fit(run, image_path, doc, max_width=None, max_height=None):
    if max_width is None or max_height is None:
        max_width, max_height = get_content_frame_size(doc)

    img_w, img_h = get_image_dimensions(image_path)
    if not img_w or not img_h:
        run.add_picture(str(image_path), width=max_width)
        return

    width_inches = img_w / 96.0
    height_inches = img_h / 96.0

    width_scale = max_width.inches / width_inches if width_inches else 1.0
    height_scale = max_height.inches / height_inches if height_inches else 1.0
    scale = min(width_scale, height_scale)

    target_width = Inches(width_inches * scale)
    target_height = Inches(height_inches * scale)
    run.add_picture(str(image_path), width=target_width, height=target_height)


def strip_md_links(text):
    return re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)


def strip_md_markup(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def is_inline_word_char(char):
    return bool(char) and (char.isalnum() or char == '_')


def can_open_inline_marker(text, idx, marker):
    marker_len = len(marker)
    prev_char = text[idx - 1] if idx > 0 else ''
    next_char = text[idx + marker_len] if idx + marker_len < len(text) else ''

    if not next_char or next_char.isspace():
        return False

    if marker.startswith('_') and is_inline_word_char(prev_char):
        return False

    return True


def can_close_inline_marker(text, idx, marker):
    marker_len = len(marker)
    prev_char = text[idx - 1] if idx > 0 else ''
    next_char = text[idx + marker_len] if idx + marker_len < len(text) else ''

    if not prev_char or prev_char.isspace():
        return False

    if marker.startswith('_') and is_inline_word_char(next_char):
        return False

    return True


def append_inline_segment(segments, text, style):
    if not text:
        return

    if segments and segments[-1][1] == style:
        segments[-1] = (segments[-1][0] + text, style)
        return

    segments.append((text, style.copy()))


def collect_inline_segments(text, start=0, end_marker=None, style=None):
    if style is None:
        style = {'bold': False, 'italic': False, 'code': False}

    segments = []
    buffer = []
    i = start

    while i < len(text):
        if end_marker and text.startswith(end_marker, i) and can_close_inline_marker(text, i, end_marker):
            append_inline_segment(segments, ''.join(buffer), style)
            return segments, i + len(end_marker), True

        if text[i] == '`':
            close_idx = text.find('`', i + 1)
            if close_idx != -1:
                append_inline_segment(segments, ''.join(buffer), style)
                buffer = []
                append_inline_segment(
                    segments,
                    text[i + 1 : close_idx],
                    {'bold': False, 'italic': False, 'code': True},
                )
                i = close_idx + 1
                continue

        matched_marker = False
        for marker, style_key in (('**', 'bold'), ('__', 'bold'), ('*', 'italic'), ('_', 'italic')):
            if not text.startswith(marker, i):
                continue
            if not can_open_inline_marker(text, i, marker):
                continue

            nested_style = style.copy()
            nested_style[style_key] = True
            nested_segments, new_idx, closed = collect_inline_segments(
                text,
                start=i + len(marker),
                end_marker=marker,
                style=nested_style,
            )
            if not closed:
                continue

            append_inline_segment(segments, ''.join(buffer), style)
            buffer = []
            segments.extend(nested_segments)
            i = new_idx
            matched_marker = True
            break

        if matched_marker:
            continue

        buffer.append(text[i])
        i += 1

    append_inline_segment(segments, ''.join(buffer), style)
    return segments, i, False


def normalize_punctuation(text):
    """Chuẩn hóa dấu câu theo NĐ30: dấu sát từ trước, space sau dấu."""
    # Xóa khoảng trắng TRƯỚC dấu câu ngắt (., , : ; ! ? )
    text = re.sub(r'\s+([.,;:!?)])', r'\1', text)
    # Xóa khoảng trắng SAU dấu mở ngoặc
    text = re.sub(r'([([{])\s+', r'\1', text)
    # Xóa khoảng trắng TRƯỚC dấu đóng ngoặc
    text = re.sub(r'\s+([)\]}])', r'\1', text)
    # Đảm bảo có đúng 1 space SAU dấu câu (bỏ qua số thập phân và URL)
    text = re.sub(r'([.,;:!?])([^\s\d.,;:!?)\]}\'"\\`*_])', r'\1 \2', text)
    return text


def resolve_media_path(md_path, asset_path):
    asset = Path(asset_path)
    if asset.is_absolute():
        return asset
    primary = (Path(md_path).resolve().parent / asset).resolve()
    if primary.exists():
        return primary
    fallback_text = asset_path.replace('../', '').replace('..\\', '').replace('./', '').replace('.\\', '')
    return (BASE / Path(fallback_text)).resolve()


def add_markdown_image(doc, md_path, image_ref):
    image_path = resolve_media_path(md_path, image_ref)
    if not image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[Khong tim thay anh: {image_ref}]')
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    max_width, max_height = get_content_frame_size(doc, height_reserve=Cm(2))
    add_picture_fit(run, image_path, doc, max_width=max_width, max_height=max_height)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table_of_contents(doc):
    p = doc.add_paragraph()
    begin_run = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    begin_run.append(fld_begin)

    instr_run = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-4" \\h \\z \\u'
    instr_run.append(instr)

    separate_run = OxmlElement('w:r')
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    separate_run.append(fld_separate)

    hint_run = OxmlElement('w:r')
    hint_text = OxmlElement('w:t')
    hint_text.text = 'Cap nhat Muc luc trong Word bang Update Field.'
    hint_run.append(hint_text)

    end_run = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    end_run.append(fld_end)

    p._p.append(begin_run)
    p._p.append(instr_run)
    p._p.append(separate_run)
    p._p.append(hint_run)
    p._p.append(end_run)


def get_heading_indent(text):
    match = re.match(r'^\s*(\d+(?:\.\d+)+)\b', text)
    if not match:
        return None
    depth = match.group(1).count('.')
    return Cm(0.5 * depth)


def add_formatted_run(para, text):
    text = strip_md_links(text)
    segments, _, _ = collect_inline_segments(text)
    for token, style in segments:
        if not token:
            continue

        run = para.add_run(token)
        if style['code']:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            continue

        if style['bold']:
            run.bold = True
        if style['italic']:
            run.italic = True


class PreviewTextExtractor(HTMLParser):
    BLOCK_TAGS = {
        'p',
        'div',
        'section',
        'article',
        'header',
        'footer',
        'aside',
        'ul',
        'ol',
        'li',
        'pre',
        'blockquote',
        'h1',
        'h2',
        'h3',
        'h4',
        'h5',
        'h6',
    }

    def __init__(self):
        super().__init__()
        self.parts = []
        self._cells_in_row = 0

    def handle_starttag(self, tag, attrs):
        if tag == 'br':
            self.parts.append('\n')
        elif tag == 'tr':
            if self.parts and not self.parts[-1].endswith('\n'):
                self.parts.append('\n')
            self._cells_in_row = 0
        elif tag in {'td', 'th'}:
            if self._cells_in_row > 0:
                self.parts.append(' | ')
            self._cells_in_row += 1

    def handle_endtag(self, tag):
        if tag == 'tr':
            self._cells_in_row = 0
            self.parts.append('\n')
        elif tag in self.BLOCK_TAGS:
            self.parts.append('\n\n')

    def handle_data(self, data):
        if not data or not data.strip():
            return
        self.parts.append(re.sub(r'\s+', ' ', data))

    def get_text(self):
        text = ''.join(self.parts)
        text = text.replace('\n| ', ' | ')
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+\n', '\n', text)
        text = re.sub(r'\n[ \t]+', '\n', text)
        lines = [line.strip() for line in text.splitlines()]
        lines = [line for line in lines if line]
        return '\n'.join(lines)


def markdown_to_html_body(md_content):
    """Chuyển Markdown sang HTML body, dùng placeholder để giữ PlantUML block."""
    lines = md_content.splitlines()
    out = []
    placeholders = {}
    idx = 0
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith('```plantuml'):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            key = f'PLANTUMLBLOCK{idx}PLACEHOLDER'
            img_url = plantuml_png_url('\n'.join(buf))
            placeholders[key] = (
                f'<p class="diagram"><img src="{quote(img_url, safe=":/?=~")}" alt="PlantUML diagram"></p>'
            )
            out.append(key)
            idx += 1
        else:
            out.append(ln)
        i += 1
    raw_md = '\n'.join(out)
    try:
        import markdown

        body = markdown.markdown(raw_md, extensions=['tables', 'fenced_code'], output_format='html')
    except ImportError:
        body = '<pre>' + html_lib.escape(raw_md) + '</pre>'
    for key, html_block in placeholders.items():
        body = body.replace('<p>' + key + '</p>', html_block)
        body = body.replace(key, html_block)
    return body


def markdown_to_preview_text(md_content):
    body = markdown_to_html_body(md_content)
    extractor = PreviewTextExtractor()
    extractor.feed(body)
    return extractor.get_text()


def inline_segments_to_preview_spans(text):
    spans = []
    segments, _, _ = collect_inline_segments(text)
    for token, style in segments:
        style_names = set()
        if style['bold']:
            style_names.add('bold')
        if style['italic']:
            style_names.add('italic')
        if style['code']:
            style_names.add('code')
        spans.append((token, style_names))
    return spans


def parse_markdown_table_row(line):
    if not line.strip().startswith('|'):
        return None
    cells = [strip_md_markup(cell.strip()) for cell in line.strip().strip('|').split('|')]
    return cells


def is_markdown_table_separator(line):
    cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
    return bool(cells) and all(c and set(c) <= {'-', ':'} for c in cells)


def markdown_to_preview_blocks(md_content):
    lines = md_content.splitlines()
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading:
            blocks.append(
                {
                    'type': 'heading',
                    'level': len(heading.group(1)),
                    'segments': inline_segments_to_preview_spans(heading.group(2).strip()),
                }
            )
            i += 1
            continue

        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('> ').strip())
                i += 1
            blocks.append(
                {
                    'type': 'quote',
                    'segments': inline_segments_to_preview_spans(' '.join(quote_lines)),
                }
            )
            continue

        if stripped.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = parse_markdown_table_row(lines[i])
                if row and not is_markdown_table_separator(lines[i]):
                    rows.append(row)
                i += 1
            if rows:
                blocks.append({'type': 'table', 'rows': rows})
            continue

        if re.match(r'^---+\s*$', stripped):
            blocks.append({'type': 'rule'})
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate:
                break
            if candidate.startswith(('>', '|', '```')) or re.match(r'^(#{1,6})\s+', candidate):
                break
            if re.match(r'^---+\s*$', candidate):
                break
            paragraph_lines.append(candidate)
            i += 1
        blocks.append(
            {
                'type': 'paragraph',
                'segments': inline_segments_to_preview_spans(' '.join(paragraph_lines)),
            }
        )

    return blocks


def configure_preview_text_tags(preview_text):
    preview_text.tag_configure('paragraph', spacing1=2, spacing3=10, lmargin1=0, lmargin2=0)
    preview_text.tag_configure('quote_block', lmargin1=18, lmargin2=18, spacing1=4, spacing3=10, foreground='#555')
    preview_text.tag_configure('table_row', lmargin1=10, lmargin2=10, spacing1=1, spacing3=1)
    preview_text.tag_configure('rule', foreground='#999', spacing1=4, spacing3=8)
    preview_text.tag_configure('bold', font=('Georgia', 10, 'bold'))
    preview_text.tag_configure('italic', font=('Georgia', 10, 'italic'))
    preview_text.tag_configure('code', font=('Consolas', 10), background='#f0f0f0', foreground='#9c2f52')
    preview_text.tag_configure('h1', font=('Georgia', 16, 'bold'), foreground='#1A3A5C', spacing1=8, spacing3=8)
    preview_text.tag_configure('h2', font=('Georgia', 14, 'bold'), foreground='#1F619E', spacing1=8, spacing3=7)
    preview_text.tag_configure('h3', font=('Georgia', 13, 'bold'), foreground='#2E86AB', spacing1=7, spacing3=6)
    preview_text.tag_configure('h4', font=('Georgia', 12, 'bold'), foreground='#449DD1', spacing1=6, spacing3=5)
    preview_text.tag_configure('h5', font=('Georgia', 11, 'bold'), foreground='#449DD1', spacing1=5, spacing3=4)
    preview_text.tag_configure('h6', font=('Georgia', 10, 'bold'), foreground='#449DD1', spacing1=4, spacing3=4)


def insert_preview_segments(preview_text, segments, block_tags):
    for token, style_names in segments:
        tags = list(block_tags) + sorted(style_names)
        preview_text.insert('end', token, tuple(tags))


def render_markdown_to_preview_widget(preview_text, md_content):
    preview_text.config(state='normal')
    preview_text.delete('1.0', 'end')

    for block in markdown_to_preview_blocks(md_content):
        if block['type'] == 'heading':
            heading_tag = f'h{min(block["level"], 6)}'
            insert_preview_segments(preview_text, block['segments'], [heading_tag])
            preview_text.insert('end', '\n\n')
        elif block['type'] == 'quote':
            preview_text.insert('end', '▌ ', ('quote_block',))
            insert_preview_segments(preview_text, block['segments'], ['quote_block'])
            preview_text.insert('end', '\n\n')
        elif block['type'] == 'table':
            for row_index, row in enumerate(block['rows']):
                row_text = ' | '.join(row)
                tags = ('table_row', 'bold') if row_index == 0 else ('table_row',)
                preview_text.insert('end', row_text, tags)
                preview_text.insert('end', '\n')
            preview_text.insert('end', '\n')
        elif block['type'] == 'rule':
            preview_text.insert('end', '─' * 42, ('rule',))
            preview_text.insert('end', '\n\n')
        else:
            insert_preview_segments(preview_text, block['segments'], ['paragraph'])
            preview_text.insert('end', '\n\n')

    preview_text.config(state='disabled')


# ── Parser MD → DOCX ─────────────────────────────────────────────────────────
def parse_and_write(doc, md_path, img_cache=IMG_CACHE):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    # Chuẩn NĐ30: justify + dãn dòng 1.5
    # space_after = 0 vì đã dùng first_line_indent để phân biệt đoạn
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Tiền xử lý: rút gọn ≥2 dòng trống liên tiếp thành tối đa 1
    compressed, blank_run = [], 0
    for _ln in lines:
        if not _ln.strip():
            blank_run += 1
            if blank_run == 1:  # chỉ giữ dòng trống đầu tiên
                compressed.append(_ln)
        else:
            blank_run = 0
            compressed.append(_ln)
    lines = compressed

    i, in_code, mermaid_block, mermaid_buf, diagram_idx, math_idx = 0, False, False, [], 0, 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code/Mermaid block
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                lang = line.strip()[3:].strip().lower()
                mermaid_block = lang in {'plantuml', 'puml'}
                mermaid_buf = []
            else:
                if mermaid_block and mermaid_buf:
                    diagram_idx += 1
                    img_path = render_plantuml('\n'.join(mermaid_buf), diagram_idx, img_cache)
                    if img_path:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run()
                        max_width, max_height = get_content_frame_size(doc, height_reserve=Cm(3))
                        add_picture_fit(run, img_path, doc, max_width=max_width, max_height=max_height)
                        cap = doc.add_paragraph(f'Bieu do {diagram_idx}')
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.runs[0].font.size = Pt(10)
                        cap.runs[0].italic = True
                        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    else:
                        p = doc.add_paragraph()
                        r = p.add_run(f'[Bieu do PlantUML {diagram_idx} - khong the render]')
                        r.italic = True
                        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                in_code = mermaid_block = False
                mermaid_buf = []
            i += 1
            continue

        if in_code:
            if mermaid_block:
                mermaid_buf.append(line)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line if line else ' ')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                set_para_shading(p, 'F4F4F4')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Bảng
        if line.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                table_rows.append(cells)
                i += 1
            data_rows = [r for r in table_rows if not all(re.match(r'^[-: ]+$', c) for c in r)]
            if not data_rows:
                continue
            max_cols = max(len(r) for r in data_rows)
            data_rows = [r + [''] * (max_cols - len(r)) for r in data_rows]
            tbl = doc.add_table(rows=len(data_rows), cols=max_cols)
            tbl.style = 'Table Grid'
            for ri, row in enumerate(data_rows):
                for ci, cell_text in enumerate(row):
                    cell = tbl.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    p = cell.paragraphs[0]
                    # Căn lề trong cell: header center, body left
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(strip_md_markup(cell_text))
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if ri == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_cell_bg(cell, '1F619E')
                    elif ri % 2 == 0:
                        set_cell_bg(cell, 'EBF4FB')
            # Khoảng cách sau bảng (spacer mỏng, không phải dòng trống đầy)
            _sp = doc.add_paragraph()
            _sp.paragraph_format.space_before = Pt(0)
            _sp.paragraph_format.space_after = Pt(6)
            _sp.add_run('').font.size = Pt(4)
            continue

        # Horizontal rule
        if re.match(r'^---+\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Centered text
        m_center = re.match(r'^\s*<center>(.*?)</center>\s*$', line)
        if m_center:
            text = m_center.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_run(p, text)
            i += 1
            continue

        # Markdown image
        m_image = re.match(r'^\s*!\[.*?\]\(([^)]+)\)\s*$', line)
        if m_image:
            add_markdown_image(doc, md_path, m_image.group(1).strip())
            i += 1
            continue

        if line.strip() == '[[PAGEBREAK]]':
            add_page_break(doc)
            i += 1
            continue

        if line.strip() == '[[TOC]]':
            add_table_of_contents(doc)
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = strip_md_markup(m.group(2))
            p = doc.add_paragraph(style=f'Heading {min(level, 4)}')
            heading_indent = get_heading_indent(text)
            if heading_indent is not None:
                p.paragraph_format.left_indent = heading_indent
            # Khoảng cách trước/sau heading — cân bằng, không quá rộng
            _space_before = [Pt(12), Pt(10), Pt(8), Pt(6), Pt(6), Pt(4)]
            p.paragraph_format.space_before = _space_before[level - 1]
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # heading luôn left
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = [COLOR_H1, COLOR_H2, COLOR_H3, COLOR_H4, COLOR_H4, COLOR_H4][level - 1]
            run.font.size = Pt([16, 14, 13, 12, 12, 11][level - 1])
            run.bold = level <= 3
            i += 1
            continue

        # Blockquote
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                quote_lines.append(lines[i].lstrip('> ').strip())
                i += 1

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_run(p, ' '.join(quote_lines))

            for run in p.runs:
                if run.font.name != 'Courier New':
                    run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Checkbox
        if re.match(r'^-\s+\[[ xX]\]', line):
            text = re.sub(r'^-\s+\[[ xX]\]\s*', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_run(p, '[  ] ' + text)
            i += 1
            continue

        # Bullet list
        m_bullet = re.match(r'^( *)[-\*]\s+(.*)', line)
        if m_bullet:
            indent_lvl = len(m_bullet.group(1)) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(indent_lvl * 1.0)
            add_formatted_run(p, m_bullet.group(2))
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, text)
            i += 1
            continue

        # Công thức toán (khối $$ ... $$)
        if line.strip() == '$$':
            # Thu thập tất cả dòng đến khi gặp $$ đóng
            math_buf = []
            i += 1
            while i < len(lines) and lines[i].rstrip('\n').strip() != '$$':
                math_buf.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # bỏ qua dòng $$ đóng
            latex_code = ' '.join(b.strip() for b in math_buf if b.strip())
            if latex_code:
                math_idx += 1
                img_path = render_latex(latex_code, math_idx, img_cache)
                if img_path:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    max_width, max_height = get_content_frame_size(doc, height_reserve=Cm(4))
                    add_picture_fit(p.add_run(), img_path, doc, max_width=max_width, max_height=max_height)
                else:
                    # Fallback: hiển thị LaTeX dạng text
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(latex_code)
                    run.font.name = 'Cambria Math'
                    run.font.size = Pt(11)
            continue

        # Công thức toán inline: $$ ... $$ trên một dòng
        if line.strip().startswith('$$') and line.strip().endswith('$$') and len(line.strip()) > 4:
            latex_code = line.strip()[2:-2].strip()
            if latex_code:
                math_idx += 1
                img_path = render_latex(latex_code, math_idx, img_cache)
                if img_path:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    max_width, max_height = get_content_frame_size(doc, height_reserve=Cm(4))
                    add_picture_fit(p.add_run(), img_path, doc, max_width=max_width, max_height=max_height)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run(latex_code).font.name = 'Cambria Math'
            i += 1
            continue

        # Dòng trống → bỏ qua hoàn toàn, không tạo paragraph
        # (space_after của đoạn văn trước đã tạo khoảng cách đủ rồi)
        if not line.strip():
            i += 1
            continue

        # Đoạn văn thường — chuẩn NĐ30: justify, lùi đầu dòng, dãn 1.5
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.27)  # lùi đầu dòng 1.27 cm
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)  # tối thiểu 6pt giữa đoạn
        add_formatted_run(p, normalize_punctuation(line.strip()))
        i += 1


def step_convert(md_out=MD_OUT, docx_out=DOCX_OUT, img_cache=IMG_CACHE):
    print('=' * 55)
    print('BƯỚC 2: Convert MD → DOCX')
    print('=' * 55)
    print(f'  Input : {md_out}')
    print(f'  Output: {docx_out}')
    print()

    os.makedirs(img_cache, exist_ok=True)
    ensure_output_docx_closed(docx_out)
    doc = Document()
    set_page_setup(doc)
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    parse_and_write(doc, md_out, img_cache)

    out = docx_out
    try:
        doc.save(out)
    except PermissionError as exc:
        raise RuntimeError(
            f'Khong the ghi file {Path(docx_out).name}. Hay dong file Word cu truoc khi build lai.'
        ) from exc

    size = os.path.getsize(out) // 1024
    print(f'\n[DONE] {out}  ({size} KB)')
    print('  => Mo file Word va dong lai neu can dung ten goc.')
    return out


def run_build_pipeline(chapters_dir=CH_DIR, md_out=MD_OUT, docx_out=DOCX_OUT, img_cache=IMG_CACHE):
    os.makedirs(img_cache, exist_ok=True)
    print('  [INFO] Dang xoa cache anh cu de lay du lieu render/API moi...')
    clear_image_cache(img_cache)
    print('=' * 55)
    print('BƯỚC 1: Ghép chapters → MD')
    print('=' * 55)
    final, chapter_files = assemble_markdown(chapters_dir, md_out)
    kb = len(final.encode('utf-8')) // 1024
    print(f'\n  => {md_out}')
    print(f'     {len(final.splitlines())} dong | {kb} KB | {len(chapter_files)} chapters\n')
    saved_path = step_convert(md_out=md_out, docx_out=docx_out, img_cache=img_cache)
    return Path(saved_path)


def launch_workflow_ui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk

    root = tk.Tk()
    root.title('NMCNPM Workflow')
    root.geometry('1400x720')  # Rộng hơn để chứa live preview panel
    root.configure(bg='#f3efe5')

    style = ttk.Style()
    style.theme_use('clam')
    style.configure('TFrame', background='#f3efe5')
    style.configure('TLabel', background='#f3efe5', foreground='#2b241b', font=('Georgia', 11))
    style.configure('Header.TLabel', background='#f3efe5', foreground='#1f3f5b', font=('Georgia', 18, 'bold'))
    style.configure('TButton', font=('Georgia', 10, 'bold'))
    style.configure('TNotebook', background='#f3efe5', borderwidth=0)
    style.configure('TNotebook.Tab', font=('Georgia', 10, 'bold'))

    container = ttk.Frame(root, padding=18)
    container.pack(fill='both', expand=True)

    ttk.Label(container, text='NMCNPM Report Workflow', style='Header.TLabel').pack(anchor='w')
    ttk.Label(
        container,
        text='Thiết lập đường dẫn và chạy các bước build, split, convert ngay trong một UI cục bộ.',
    ).pack(anchor='w', pady=(4, 14))

    notebook = ttk.Notebook(container)
    notebook.pack(fill='both', expand=True)

    log_box = tk.Text(container, height=10, bg='#fffdf8', fg='#2b241b', font=('Consolas', 10))
    log_box.pack(fill='x', pady=(14, 0))

    def log(message):
        log_box.insert('end', message + '\n')
        log_box.see('end')
        root.update_idletasks()

    def add_path_row(parent, label, variable, browse_kind, filetypes=None):
        row = ttk.Frame(parent)
        row.pack(fill='x', pady=6)
        ttk.Label(row, text=label, width=18).pack(side='left')
        entry = ttk.Entry(row, textvariable=variable)
        entry.pack(side='left', fill='x', expand=True, padx=(0, 8))

        def choose():
            if browse_kind == 'dir':
                value = filedialog.askdirectory(initialdir=str(BASE))
            elif browse_kind == 'save':
                value = filedialog.asksaveasfilename(initialdir=str(BASE), filetypes=filetypes)
            else:
                value = filedialog.askopenfilename(initialdir=str(BASE), filetypes=filetypes)
            if value:
                variable.set(value)

        ttk.Button(row, text='Browse', command=choose).pack(side='left')
        return entry

    editor_tab = ttk.Frame(notebook, padding=14)
    build_tab = ttk.Frame(notebook, padding=14)
    split_tab = ttk.Frame(notebook, padding=14)
    convert_tab = ttk.Frame(notebook, padding=14)
    notebook.add(editor_tab, text='Editor')
    notebook.add(build_tab, text='Build DOCX')
    notebook.add(split_tab, text='Split Chapters')
    notebook.add(convert_tab, text='DOCX → MD')

    # ── Kiểm tra tkinterweb ──────────────────────────────────────────────────
    _htmlframe_error = None
    try:
        from tkinterweb import HtmlFrame as _HtmlFrame

        _has_htmlframe = True
    except Exception as exc:
        _has_htmlframe = False
        _htmlframe_error = str(exc)

    # --- EDITOR TAB ---
    # Bố cục: [Listbox file | Editor text | Preview panel]
    editor_paned = ttk.PanedWindow(editor_tab, orient=tk.HORIZONTAL)
    editor_paned.pack(fill='both', expand=True)

    # Cột 1: Danh sách file
    file_list_frame = ttk.Frame(editor_paned)
    editor_paned.add(file_list_frame, weight=1)

    # Cột 2: Editor text
    editor_frame = ttk.Frame(editor_paned)
    editor_paned.add(editor_frame, weight=3)

    # Cột 3: Live preview
    preview_frame = ttk.Frame(editor_paned)
    editor_paned.add(preview_frame, weight=3)

    # -- Listbox --
    file_listbox = tk.Listbox(file_list_frame, font=('Consolas', 10))
    file_listbox.pack(side='left', fill='both', expand=True)
    scrollbar = ttk.Scrollbar(file_list_frame, orient='vertical', command=file_listbox.yview)
    scrollbar.pack(side='right', fill='y')
    file_listbox.config(yscrollcommand=scrollbar.set)

    # -- Editor --
    editor_text = tk.Text(editor_frame, font=('Consolas', 11), wrap='word', undo=True)
    editor_text.pack(fill='both', expand=True, pady=(0, 5))

    btn_frame = ttk.Frame(editor_frame)
    btn_frame.pack(fill='x')

    # -- Live preview widget --
    ttk.Label(preview_frame, text='Live Preview', font=('Georgia', 10, 'bold')).pack(anchor='w', padx=6, pady=(0, 4))
    if _has_htmlframe:
        # Dùng HtmlFrame (tkinterweb) — render HTML thực sự
        live_preview = _HtmlFrame(preview_frame, messages_enabled=False)
        live_preview.pack(fill='both', expand=True)

        def _update_preview_widget(html_content):
            live_preview.load_html(html_content)
    else:
        # Fallback: tk.Text hiển thị nội dung thô (chỉ dùng khi tkinterweb chưa cài)
        preview_text = tk.Text(
            preview_frame, font=('Consolas', 10), wrap='word', bg='#fffdf8', fg='#2b241b', state='disabled'
        )
        preview_text.pack(fill='both', expand=True)
        configure_preview_text_tags(preview_text)
        ttk.Label(
            preview_frame,
            text=(
                'Đang dùng fallback preview trong ứng dụng.\n'
                'Muốn render HTML đầy đủ: cài `tkinterweb` vào đúng Python đang chạy UI.\n'
                f'Python hiện tại: {sys.executable}\n'
                + (f'Lý do không tải được HtmlFrame: {_htmlframe_error}' if _htmlframe_error else '')
            ),
            foreground='#888',
            font=('Consolas', 9),
        ).pack(anchor='w', padx=6)

        def _update_preview_widget(html_content):
            # Fallback: render markdown có style ngay trong Text widget
            render_markdown_to_preview_widget(preview_text, editor_text.get(1.0, tk.END))

    # -- Debounce timer cho live preview --
    _preview_after_id = [None]  # dùng list để có thể thay đổi trong closure

    def _build_preview_html():
        """Tạo HTML đầy đủ từ nội dung editor hiện tại."""
        content = editor_text.get(1.0, tk.END)
        filepath = current_file.get()
        body = markdown_to_html_body(content)
        fname = os.path.basename(filepath) if filepath else 'Preview'
        css = """
body { font-family: 'Segoe UI', Arial, sans-serif; padding: 20px 32px;
       line-height: 1.7; color: #222; max-width: 860px; margin: auto; }
h1 { color: #1A3A5C; font-size: 1.7em; border-bottom: 2px solid #1F619E; padding-bottom: 5px; }
h2 { color: #1F619E; font-size: 1.35em; }
h3 { color: #2E86AB; font-size: 1.15em; }
h4 { color: #449DD1; }
pre { background:#f4f4f4; padding:10px; border-radius:5px; overflow-x:auto;
      font-family: Consolas, monospace; font-size: 0.9em; }
code { background:#f0f0f0; padding:2px 4px; border-radius:3px; font-size:0.9em; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; }
th { background-color: #1F619E; color: white; padding: 7px 10px; text-align: left; }
td { border: 1px solid #d0d0d0; padding: 6px 10px; }
tr:nth-child(even) td { background: #EBF4FB; }
blockquote { border-left: 4px solid #449DD1; margin-left:0; padding-left:14px; color:#555; }
img { max-width: 100%; height: auto; display: block; margin: 8px auto; }
.diagram { text-align: center; }
"""
        return f"""<!DOCTYPE html>
<html lang="vi">
<head><meta charset="utf-8"><title>{fname}</title><style>{css}</style></head>
<body>{body}</body>
</html>"""

    def _do_live_preview():
        """Được gọi sau debounce — render và cập nhật preview."""
        try:
            html_content = _build_preview_html()
            _update_preview_widget(html_content)
        except Exception:
            pass  # Không để lỗi render crash UI
        _preview_after_id[0] = None

    def _on_editor_change(event=None):
        """Debounce 800ms: chỉ render sau khi ngừng gõ."""
        if _preview_after_id[0] is not None:
            root.after_cancel(_preview_after_id[0])
        _preview_after_id[0] = root.after(800, _do_live_preview)

    # Gắn sự kiện gõ phím vào editor
    editor_text.bind('<KeyRelease>', _on_editor_change)

    current_file = tk.StringVar()

    def load_files():
        file_listbox.delete(0, tk.END)
        for f in collect_chapter_files(CH_DIR):
            file_listbox.insert(tk.END, os.path.basename(f))

    def on_file_select(event):
        if not file_listbox.curselection():
            return
        idx = file_listbox.curselection()[0]
        filename = file_listbox.get(idx)
        filepath = os.path.join(CH_DIR, filename)
        current_file.set(filepath)
        try:
            with open(filepath, encoding='utf-8') as f:
                editor_text.delete(1.0, tk.END)
                editor_text.insert(tk.END, f.read())
            # Render ngay khi mở file mới
            _do_live_preview()
        except Exception as e:
            log(f'Lỗi đọc file: {e}')

    file_listbox.bind('<<ListboxSelect>>', on_file_select)

    def save_file():
        filepath = current_file.get()
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(editor_text.get(1.0, tk.END).rstrip('\n') + '\n')
                log(f'Đã lưu: {os.path.basename(filepath)}')
            except Exception as e:
                log(f'Lỗi lưu file: {e}')

    def preview_file():
        filepath = current_file.get()
        if not filepath:
            messagebox.showwarning('Cảnh báo', 'Vui lòng chọn file để preview!')
            return
        import tempfile
        import webbrowser

        content = editor_text.get(1.0, tk.END)
        body = markdown_to_html_body(content)
        css = """
body { font-family: 'Segoe UI', Arial, sans-serif; padding: 28px 48px;
       line-height: 1.7; color: #222; max-width: 860px; margin: auto; }
h1 { color: #1A3A5C; font-size: 1.8em; border-bottom: 2px solid #1F619E; padding-bottom: 6px; }
h2 { color: #1F619E; font-size: 1.4em; }
h3 { color: #2E86AB; font-size: 1.2em; }
h4 { color: #449DD1; }
pre { background:#f4f4f4; padding:12px; border-radius:6px; overflow-x:auto;
      font-family: Consolas, monospace; font-size: 0.92em; }
code { background:#f0f0f0; padding:2px 5px; border-radius:3px; font-size:0.92em; }
table { border-collapse: collapse; width: 100%; margin: 14px 0; }
th { background-color: #1F619E; color: white; padding: 8px 12px; text-align: left; }
td { border: 1px solid #d0d0d0; padding: 7px 12px; }
tr:nth-child(even) td { background: #EBF4FB; }
blockquote { border-left: 4px solid #449DD1; margin-left:0; padding-left:16px; color:#555; }
img { max-width: 100%; }
.diagram { text-align: center; }
.diagram img { max-width: 100%; height: auto; }
"""
        html_content = f"""<!DOCTYPE html>
<html lang="vi">
<head>
  <meta charset="utf-8">
  <title>Preview — {os.path.basename(filepath)}</title>
  <style>{css}</style>
</head>
<body>
{body}
</body>
</html>"""
        fd, temp_path = tempfile.mkstemp(suffix='.html')
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            f.write(html_content)
        webbrowser.open('file:///' + temp_path.replace('\\', '/'))
        log(f'Preview: {os.path.basename(filepath)}')

    def docx_to_chapters():
        docx_path = filedialog.askopenfilename(initialdir=str(BASE), filetypes=[('Word', '*.docx')])
        if not docx_path:
            return
        log('Đang tách DOCX -> Markdown...')
        try:
            tmp_md = str(BASE / 'temp_split.md')
            convert_docx_to_markdown(docx_path, tmp_md, str(BASE / 'extracted_media' / 'temp_split'))
            split_chapters.write_chapter_files(tmp_md, CH_DIR)
            if os.path.exists(tmp_md):
                os.remove(tmp_md)
            log('Tách thành công!')
            load_files()
            messagebox.showinfo('Thành công', 'Đã tách DOCX thành các chapters.')
        except Exception as e:
            log(f'Lỗi: {e}')

    ttk.Button(btn_frame, text='Lưu (Save)', command=save_file).pack(side='left', padx=5)
    ttk.Button(btn_frame, text='Mở Preview (Browser)', command=preview_file).pack(side='left', padx=5)
    ttk.Button(btn_frame, text='Tải lại danh sách', command=load_files).pack(side='left', padx=5)
    ttk.Button(btn_frame, text='Tách DOCX -> Chapters', command=docx_to_chapters).pack(side='right', padx=5)

    load_files()

    build_chapters = tk.StringVar(value=CH_DIR)
    build_md = tk.StringVar(value=MD_OUT)
    build_docx = tk.StringVar(value=DOCX_OUT)
    build_cache = tk.StringVar(value=IMG_CACHE)

    add_path_row(build_tab, 'Chapters dir', build_chapters, 'dir')
    add_path_row(build_tab, 'Output MD', build_md, 'save', [('Markdown', '*.md')])
    add_path_row(build_tab, 'Output DOCX', build_docx, 'save', [('Word Document', '*.docx')])
    add_path_row(build_tab, 'Diagram cache', build_cache, 'dir')

    build_res_frame = ttk.Frame(build_tab)
    build_res_frame.pack(anchor='w', pady=(10, 0))

    def do_build():
        res = run_ui_action(
            root,
            log,
            lambda: run_build_pipeline(
                chapters_dir=build_chapters.get(),
                md_out=build_md.get(),
                docx_out=build_docx.get(),
                img_cache=build_cache.get(),
            ),
            'Build hoàn tất.',
        )
        if res:
            btn_open_file.pack(side='left', padx=5)
            btn_open_folder.pack(side='left', padx=5)

    ttk.Button(build_tab, text='Run Build Pipeline', command=do_build).pack(anchor='w', pady=(14, 0))

    btn_open_file = ttk.Button(build_res_frame, text='Mở file Word', command=lambda: os.startfile(build_docx.get()))
    btn_open_folder = ttk.Button(
        build_res_frame,
        text='Mở thư mục chứa',
        command=lambda: os.startfile(os.path.dirname(os.path.abspath(build_docx.get()))),
    )

    split_source = tk.StringVar(value=MD_OUT)
    split_output = tk.StringVar(value=str(BASE / 'chapters'))
    add_path_row(split_tab, 'Source MD', split_source, 'open', [('Markdown', '*.md')])
    add_path_row(split_tab, 'Output dir', split_output, 'dir')
    ttk.Button(
        split_tab,
        text='Split Markdown',
        command=lambda: run_ui_action(
            root,
            log,
            lambda: split_chapters.write_chapter_files(split_source.get(), split_output.get()),
            'Tách chapter hoàn tất.',
        ),
    ).pack(anchor='w', pady=(14, 0))

    convert_input = tk.StringVar(value=DOCX_OUT)
    convert_output = tk.StringVar(value=str(BASE / 'Bao_Cao_Tieu_Luan_NMCNPM_from_docx.md'))
    convert_media = tk.StringVar(value=str(BASE / 'extracted_media' / 'Bao_Cao_Tieu_Luan_NMCNPM'))
    add_path_row(convert_tab, 'Input DOCX', convert_input, 'open', [('Word Document', '*.docx')])
    add_path_row(convert_tab, 'Output MD', convert_output, 'save', [('Markdown', '*.md')])
    add_path_row(convert_tab, 'Media dir', convert_media, 'dir')
    ttk.Button(
        convert_tab,
        text='Convert DOCX to Markdown',
        command=lambda: run_ui_action(
            root,
            log,
            lambda: convert_docx_to_markdown(convert_input.get(), convert_output.get(), convert_media.get()),
            'Convert DOCX → Markdown hoàn tất.',
        ),
    ).pack(anchor='w', pady=(14, 0))

    messagebox.showinfo(
        'NMCNPM Workflow',
        'UI đã sẵn sàng. Bạn có thể thiết lập đường dẫn file/folder và chạy từng bước ngay tại đây.',
    )
    root.mainloop()


def run_ui_action(root, log, action, success_message):
    try:
        log('---')
        result = action()
        log(f'OK: {result}')
        from tkinter import messagebox

        messagebox.showinfo('NMCNPM Workflow', success_message)
        return result
    except Exception as exc:
        log(f'ERROR: {exc}')
        from tkinter import messagebox

        messagebox.showerror('NMCNPM Workflow', str(exc))


def parse_args():
    parser = argparse.ArgumentParser(description='NMCNPM report workflow')
    parser.add_argument('--ui', action='store_true', help='Launch local workflow UI')
    parser.add_argument('--chapters-dir', default=CH_DIR)
    parser.add_argument('--md-out', default=MD_OUT)
    parser.add_argument('--docx-out', default=DOCX_OUT)
    parser.add_argument('--img-cache', default=IMG_CACHE)
    return parser.parse_args()


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    args = parse_args()
    if args.ui:
        launch_workflow_ui()
    else:
        run_build_pipeline(
            chapters_dir=args.chapters_dir,
            md_out=args.md_out,
            docx_out=args.docx_out,
            img_cache=args.img_cache,
        )
