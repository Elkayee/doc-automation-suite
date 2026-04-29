"""
gen_bien_ban.py — Tạo 3 biên bản họp nhóm 6 theo format mẫu Lần 1.
Chạy: python gen_bien_ban.py
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.stdout.reconfigure(encoding='utf-8')

# ── Thông tin nhóm ──────────────────────────────────────────────────────────
THANH_VIEN = [
    'Vũ Tất Thành ( Trưởng nhóm)',
    'Nguyễn Viết Tùng',
    'Nguyễn Quang Đạo',
    'Tạ Bảo Anh Ngọc',
    'Ngô Thị Hồng Nhung',
]


# ── Helper: tạo file DOCX theo đúng cấu trúc mẫu ──────────────────────────
def tao_bien_ban(ten_file: str, lan: int, ngay: str, gio: str, noi_dung: list, ket_luan: list):
    doc = Document()

    # Cài đặt trang: A4, margin chuẩn
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.5)

    # Đặt font mặc định
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # ── Tiêu đề ──────────────────────────────────────────────────────────────
    tieu_de = doc.add_paragraph()
    tieu_de.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tieu_de.add_run(f'BIÊN BẢN HỌP LẦN {lan}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    tieu_de.paragraph_format.space_after = Pt(12)

    # ── 1. Thành phần ────────────────────────────────────────────────────────
    them_muc(doc, '1. THÀNH PHẦN')
    them_dong(doc, 'Các thành viên:')
    for tv in THANH_VIEN:
        them_dong(doc, tv)

    # ── 2. Ngày họp ──────────────────────────────────────────────────────────
    them_muc(doc, '2. NGÀY HỌP')
    them_bullet(doc, f'Ngày: {ngay}')
    them_bullet(doc, f'Thời gian: {gio}')
    them_bullet(doc, 'Hình thức: Online qua Google Meet')

    # ── 3. Nội dung ──────────────────────────────────────────────────────────
    them_muc(doc, '3. NỘI DUNG')
    for dong in noi_dung:
        # Phân loại: dòng '+' là sub-bullet, còn lại là bullet chính
        if dong.startswith('+'):
            them_sub_bullet(doc, dong[1:].strip())
        else:
            them_bullet(doc, dong)

    # ── 4. Hình ảnh / Videos ─────────────────────────────────────────────────
    them_muc(doc, '4. HÌNH ẢNH / VIDEOS')
    them_dong(doc, '(Hình ảnh chụp màn hình buổi họp — đính kèm riêng)')

    # ── 5. Kết luận ──────────────────────────────────────────────────────────
    them_muc(doc, '5. KẾT LUẬN')
    for dong in ket_luan:
        if dong.startswith('+'):
            them_sub_bullet(doc, dong[1:].strip())
        else:
            them_bullet(doc, dong)

    # Lưu file
    out = Path(__file__).resolve().parent / ten_file
    doc.save(str(out))
    print(f'[OK] {out}')


def them_muc(doc, text):
    """Dòng tiêu đề mục (bold)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)


def them_dong(doc, text):
    """Dòng text thường."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)


def them_bullet(doc, text):
    """Dòng dấu '-'."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'- {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)


def them_sub_bullet(doc, text):
    """Dòng dấu '+'."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'+ {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)


# ════════════════════════════════════════════════════════════════════════════
# BIÊN BẢN LẦN 1 — Giới thiệu đề tài & thống nhất cách làm
# ════════════════════════════════════════════════════════════════════════════
noi_dung_1 = [
    'Giới thiệu tổng quan đề tài: Hệ thống quản lý chuỗi quán café (Café Chain Management System).',
    'Thảo luận và thống nhất hướng tiếp cận:',
    '+ Hệ thống phục vụ 3 nhóm người dùng chính: Quản lý, Thu ngân và Nhân viên.',
    '+ Tập trung vào 4 phân hệ nghiệp vụ: Bán hàng, Thực đơn, Kho nguyên liệu, Nhân sự & Ca làm việc.',
    'Thống nhất cách vẽ biểu đồ:',
    '+ Biểu đồ Use Case tổng quát: vẽ bằng PlantUML, thể hiện đủ Actor và Use Case chính.',
    '+ Biểu đồ Use Case chi tiết: dùng PlantUML, bổ sung include/extend khi cần.',
    '+ Biểu đồ hoạt động (Activity Diagram): dùng PlantUML, thể hiện rõ luồng chính và luồng ngoại lệ, có swimlane phân biệt Actor và Hệ thống.',
    '+ Biểu đồ lớp thực thể (Entity Class Diagram): dùng PlantUML, thể hiện thuộc tính, kiểu dữ liệu và quan hệ giữa các lớp.',
    '+ Quy tắc đặt tên bảng CSDL: tiếng Việt không dấu, viết thường, phân cách bằng dấu gạch dưới (ví dụ: nhan_vien, hoa_don).',
    'Thống nhất cách viết báo cáo:',
    '+ Bố cục: Phần 1 (nội dung chung cả nhóm) và Phần 2 (nội dung riêng từng thành viên).',
    '+ Font chữ: Times New Roman, cỡ 13pt, dãn dòng 1.5.',
    '+ Lề trang: Trái 3 cm, Phải 2 cm, Trên 2.5 cm, Dưới 2.5 cm.',
    '+ Tiêu đề mục: in đậm, viết hoa chữ đầu, có đánh số thứ tự.',
    '+ Bảng biểu: hàng tiêu đề in đậm, căn giữa; dữ liệu căn trái.',
    '+ Mỗi biểu đồ phải có chú thích (caption) bên dưới.',
    '+ Đặc tả Use Case trình bày dạng bảng: Mã UC, Tên UC, Actor, Điều kiện tiên quyết, Luồng chính, Luồng ngoại lệ.',
    'Phân chia công việc ban đầu:',
    '+ Vũ Tất Thành: Phụ trách phân hệ bán hàng, đơn hàng và quản lý bàn (UC02, UC03).',
    '+ Tạ Bảo Anh Ngọc: Phụ trách quản lý menu và công thức pha chế (UC01).',
    '+ Nguyễn Quang Đạo: Phụ trách quản lý kho và nguyên liệu (UC05).',
    '+ Nguyễn Viết Tùng: Phụ trách quản lý ca làm việc, chấm công và phân quyền (UC04, UC07).',
    '+ Ngô Thị Hồng Nhung: Phụ trách báo cáo doanh thu, chi phí và quản lý cửa hàng (UC06).',
    '+ Nhóm cùng thực hiện: Mô tả bài toán, biểu đồ UC tổng quát, biểu đồ lớp thực thể.',
]
ket_luan_1 = [
    'Thống nhất tên đề tài và hướng tiếp cận của nhóm.',
    'Đã thống nhất quy chuẩn vẽ biểu đồ (PlantUML) và cách viết báo cáo.',
    'Đã phân chia công việc ban đầu cho từng thành viên.',
    'Bước tiếp theo:',
    '+ Mỗi thành viên tìm hiểu Use Case của mình và chuẩn bị phác thảo trước buổi họp Lần 2.',
    '+ Nhóm phụ trách bắt đầu phác thảo mô tả bài toán và danh sách Actor.',
]

noi_dung_2 = [
    'Review kết quả sau buổi họp Lần 1 — phân chia ban đầu đã rõ ràng.',
    'Thống nhất chính thức đề tài: Hệ thống quản lý chuỗi quán café (Café Chain Management System).',
    'Xác định các Actor của hệ thống:',
    '+ Quản lý (Admin): Cấu hình hệ thống, quản lý nhân sự, xem báo cáo.',
    '+ Thu ngân: Tạo đơn hàng, thanh toán, in hóa đơn, quản lý bàn.',
    '+ Nhân viên pha chế / phục vụ: Nhận đơn, cập nhật trạng thái đơn.',
    '+ Nhân viên theo ca: Xem lịch làm, chấm công vào/ra ca.',
    'Thảo luận và xác định Use Case tổng quát (6 UC chính):',
    '+ UC01 — Quản lý menu & công thức pha chế (Tạ Bảo Anh Ngọc)',
    '+ UC02, UC03 — Quản lý đơn hàng, bàn & thanh toán (Vũ Tất Thành)',
    '+ UC04 — Quản lý ca làm việc & chấm công (Nguyễn Viết Tùng)',
    '+ UC05 — Quản lý kho & nguyên liệu (Nguyễn Quang Đạo)',
    '+ UC06 — Báo cáo doanh thu & chi phí (Ngô Thị Hồng Nhung)',
    '+ UC07 — Quản lý tài khoản & phân quyền nhân sự (Nguyễn Viết Tùng)',
    'Xác định yêu cầu phi chức năng:',
    '+ Hệ thống phải hoạt động ổn định trong giờ cao điểm, thao tác nhanh, ít bước.',
    '+ Dữ liệu chấm công, doanh thu và tồn kho phải có khả năng đối soát lại khi cần.',
    '+ Chỉ người có quyền phù hợp mới được thay đổi dữ liệu quan trọng.',
    'Phân chia nhiệm vụ tiếp theo:',
    '+ Mỗi thành viên tự xây dựng biểu đồ UC chi tiết cho UC của mình.',
    '+ Nhóm phụ trách phần chung: xây dựng biểu đồ UC tổng quát và biểu đồ lớp thực thể.',
]
ket_luan_2 = [
    'Thống nhất tên đề tài và danh sách Use Case chính thức.',
    'Đã xác định đầy đủ Actor và phân công UC cho từng thành viên.',
    'Bước tiếp theo:',
    '+ Mỗi thành viên phác thảo biểu đồ UC chi tiết (UC diagram + đặc tả) trước buổi họp Lần 3.',
    '+ Nhóm thống nhất phác thảo ERD và biểu đồ lớp thực thể.',
]

# ════════════════════════════════════════════════════════════════════════════
# BIÊN BẢN LẦN 3 — Review thiết kế CSDL & biểu đồ lớp thực thể
# ════════════════════════════════════════════════════════════════════════════
noi_dung_3 = [
    'Từng thành viên trình bày phác thảo biểu đồ UC chi tiết của UC mình phụ trách.',
    'Thảo luận và góp ý chéo về luồng chính, luồng ngoại lệ của từng UC:',
    '+ UC04 (Nguyễn Viết Tùng): Bổ sung ngoại lệ vào ca trùng, vào ca muộn >15 phút.',
    '+ UC05 (Nguyễn Quang Đạo): Bổ sung luồng cảnh báo tồn kho dưới ngưỡng tối thiểu.',
    '+ UC06 (Ngô Thị Hồng Nhung): Làm rõ phân biệt báo cáo doanh thu theo ngày / tháng / chi nhánh.',
    '+ UC01 (Tạ Bảo Anh Ngọc): Bổ sung ràng buộc công thức phải có ít nhất 1 nguyên liệu.',
    'Review phác thảo biểu đồ lớp thực thể (Entity Class Diagram):',
    '+ Xác định các thực thể chính: NhanVien, TaiKhoan, Shift, ShiftAssignment, Attendance, DoUong, ToNhom, Topping, CongThuc, NguyenLieu, TonKho, HoaDon, HoaDonChiTiet, Ban, KhuVuc, CuaHang, BaoCaoDoanhThu, ChiPhi.',
    '+ Thảo luận về quan hệ giữa các thực thể: 1-N, N-N và bảng trung gian.',
    '+ Thống nhất quy tắc đặt tên bảng: tiếng Việt không dấu, viết thường, phân cách bằng dấu gạch dưới.',
    'Thống nhất ràng buộc nghiệp vụ quan trọng:',
    '+ BR-GEN-03: Tồn kho phải giảm theo công thức pha chế khi đơn hàng hoàn tất.',
    '+ BR-GEN-04: Một nhân viên không được phân công 2 ca chồng chéo trong cùng ngày.',
    '+ BR-GEN-05: Dữ liệu chấm công phải ghi nhận đúng thời điểm, không trùng lặp.',
    'Phân công hoàn thiện nội dung:',
    '+ Mỗi thành viên hoàn thiện biểu đồ UC chi tiết, đặc tả UC và biểu đồ hoạt động.',
    '+ Nhóm hoàn thiện biểu đồ lớp thực thể chung và mô tả từng thực thể.',
]
ket_luan_3 = [
    'Đã review và góp ý xong biểu đồ UC của tất cả thành viên.',
    'Thống nhất cấu trúc CSDL và danh sách thực thể chính thức.',
    'Bước tiếp theo:',
    '+ Hoàn thiện phần nội dung cá nhân (UC chi tiết, đặc tả, biểu đồ hoạt động).',
    '+ Họp Lần 4 để review tổng thể và ghép báo cáo hoàn chỉnh.',
]

# ════════════════════════════════════════════════════════════════════════════
# BIÊN BẢN LẦN 4 — Review tiến độ & thống nhất báo cáo cuối
# ════════════════════════════════════════════════════════════════════════════
noi_dung_4 = [
    'Từng thành viên báo cáo tiến độ hoàn thiện nội dung cá nhân:',
    '+ Nguyễn Viết Tùng (UC04, UC07): Đã hoàn thiện biểu đồ UC, đặc tả và biểu đồ hoạt động.',
    '+ Nguyễn Quang Đạo (UC05): Đã hoàn thiện biểu đồ UC, đang bổ sung biểu đồ hoạt động.',
    '+ Tạ Bảo Anh Ngọc (UC01): Đã hoàn thiện biểu đồ UC và đặc tả.',
    '+ Ngô Thị Hồng Nhung (UC06): Đã hoàn thiện đặc tả, đang hoàn thiện biểu đồ hoạt động.',
    '+ Vũ Tất Thành (UC02, UC03): Đã hoàn thiện biểu đồ UC và biểu đồ hoạt động.',
    'Review nội dung chung của nhóm:',
    '+ Kiểm tra phần mô tả bài toán: đảm bảo viết đúng theo ngôn ngữ tự nhiên như góc nhìn của khách hàng.',
    '+ Kiểm tra biểu đồ UC tổng quát: đã thể hiện đủ 7 Use Case chính và các Actor.',
    '+ Kiểm tra biểu đồ lớp thực thể: thống nhất quan hệ giữa các thực thể, bổ sung multiplicity.',
    'Thống nhất format và bố cục báo cáo cuối:',
    '+ Phần 1 (nội dung chung): mô tả bài toán → đặc tả yêu cầu → biểu đồ UC tổng quát → biểu đồ lớp thực thể.',
    '+ Phần 2 (nội dung cá nhân): biểu đồ UC chi tiết → đặc tả UC → biểu đồ hoạt động.',
    '+ Font: Times New Roman 13pt, dãn dòng 1.5, margin: Trái 3cm / Phải 2cm / Trên-Dưới 2.5cm.',
    'Thảo luận về hạn nộp và phân công ghép báo cáo:',
    '+ Hạn nộp chính thức: trước 23h ngày 25/04/2026 qua hệ thống LMS của nhà trường.',
    '+ Vũ Tất Thành phụ trách ghép toàn bộ nội dung và xuất file DOCX cuối.',
    '+ Mỗi thành viên gửi nội dung cá nhân hoàn chỉnh trước 22h ngày 25/04/2026.',
]
ket_luan_4 = [
    'Tất cả thành viên đã hoàn thành hoặc sắp hoàn thành nội dung cá nhân.',
    'Thống nhất bố cục và format file báo cáo cuối.',
    'Bước tiếp theo:',
    '+ Từng thành viên gửi nội dung hoàn chỉnh trước 22h ngày 25/04/2026.',
    '+ Trưởng nhóm ghép báo cáo, kiểm tra lần cuối và nộp lên LMS trước 23h ngày 25/04/2026.',
]

# ── Tạo 4 file ──────────────────────────────────────────────────────────────
tao_bien_ban('Group_6_Bien_ban_hop_Lan1.docx', 1, '06/04/2026', '22h00~23h30', noi_dung_1, ket_luan_1)
tao_bien_ban('Group_6_Bien_ban_hop_Lan2.docx', 2, '12/04/2026', '21h00~22h30', noi_dung_2, ket_luan_2)
tao_bien_ban('Group_6_Bien_ban_hop_Lan3.docx', 3, '18/04/2026', '21h30~23h00', noi_dung_3, ket_luan_3)
tao_bien_ban('Group_6_Bien_ban_hop_Lan4.docx', 4, '23/04/2026', '22h00~23h30', noi_dung_4, ket_luan_4)
print('\n[DONE] Đã tạo xong 4 biên bản họp.')
