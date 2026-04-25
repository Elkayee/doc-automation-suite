"""
Script chuyển đổi MD sang DOCX theo format template
Sử dụng python-docx để tạo DOCX với format chuẩn văn bản hành chính
"""

import sys
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_run_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    # Set font for Asian characters
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph_with_format(doc, text, alignment='left', font_size=12, bold=False, italic=False, spacing_after=6):
    """Add a paragraph with specific formatting"""
    para = doc.add_paragraph()

    # Set alignment
    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set spacing
    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = 1.15

    # Add text with formatting
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold, italic=italic)

    return para

def parse_md_line(line):
    """Parse markdown formatting from a line"""
    # Remove ** for bold markers
    text = line.strip()
    is_bold = text.startswith('**') and '**' in text[2:]
    is_italic = text.startswith('*') and not text.startswith('**')

    # Clean markdown syntax
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Remove italic markers

    return text, is_bold, is_italic

def add_mixed_paragraph(doc, text, alignment='justify', base_size=12, spacing_after=6):
    """Add paragraph with mixed bold/normal text"""
    para = doc.add_paragraph()

    if alignment == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    para.paragraph_format.space_after = Pt(spacing_after)
    para.paragraph_format.line_spacing = 1.15

    # Parse bold sections **text** and italic *text*
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, font_size=base_size, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            set_run_font(run, font_size=base_size, italic=True)
        else:
            run = para.add_run(part)
            set_run_font(run, font_size=base_size)

    return para

def convert_md_to_docx(md_file, output_file):
    """Convert MD file to DOCX with proper formatting"""

    # Read MD content
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create new document
    doc = Document()

    # Set up page margins (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2)

    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Section headers (#, ##, ###)
        if line.startswith('### '):
            add_paragraph_with_format(doc, line[4:], alignment='left', font_size=12, bold=True, spacing_after=6)
            i += 1
            continue
        if line.startswith('## '):
            add_paragraph_with_format(doc, line[3:], alignment='left', font_size=13, bold=True, spacing_after=8)
            i += 1
            continue
        if line.startswith('# '):
            add_paragraph_with_format(doc, line[2:], alignment='center', font_size=14, bold=True, spacing_after=12)
            i += 1
            continue

        # List items (- or numbered)
        if line.startswith('- ') or line.startswith('+ '):
            clean_text = line[2:].strip()
            para = add_mixed_paragraph(doc, '• ' + clean_text, alignment='justify', base_size=12)
            para.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue

        # Regular paragraph with mixed formatting
        add_mixed_paragraph(doc, line, alignment='justify', base_size=12)
        i += 1

    # Save document
    doc.save(output_file)
    print(f"Created: {output_file}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python convert_md_to_docx.py input.md output.docx")
        sys.exit(1)

    convert_md_to_docx(sys.argv[1], sys.argv[2])
