import shutil
import unittest
import zipfile
from pathlib import Path

import yaml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from src.core.docx_builder import DocxBuilder
from src.core.docx_helpers import DocxHelpers

PROJECT_ROOT = Path(__file__).resolve().parent.parent


class DocxBuilderListMarkerTests(unittest.TestCase):
    def test_build_preserves_literal_bullet_marker_from_markdown(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_list_markers'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Test',
                        'description': '',
                        'type': 'report',
                        'docx_template': 'template.docx',
                        'required_files': ['Ch01_Test.md'],
                        'settings': {
                            'chapter_settings': {
                                'Ch01_Test.md': {
                                    'list_markers_by_level': ['-', '+', '*'],
                                }
                            }
                        },
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: Ch01_Test.md -->\n\n+ Muc cap 1\n  * Muc cap 2\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)

            paragraph_texts = [paragraph.text for paragraph in builder.doc.paragraphs if paragraph.text.strip()]
            self.assertIn('+ Muc cap 1', paragraph_texts)
            self.assertIn('* Muc cap 2', paragraph_texts)
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)

    def test_build_overrides_exact_line_spacing_for_image_paragraphs(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_image_spacing'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'
            output_path = workspace / 'out.docx'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Test',
                        'description': '',
                        'type': 'report',
                        'docx_template': 'template.docx',
                        'required_files': ['Ch01_Test.md'],
                        'settings': {
                            'line_spacing_mode': 'exactly',
                            'line_spacing_value': 16.0,
                        },
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: Ch01_Test.md -->\n\n'
                'Doan truoc.\n\n'
                f'![Image]({PROJECT_ROOT.as_posix()}/test_extracted.png)\n\n'
                'Doan sau.\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)
            builder.save(output_path)

            with zipfile.ZipFile(output_path, 'r') as archive:
                document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')

            image_paragraph = next(
                (paragraph_xml for paragraph_xml in document_xml.split('<w:p') if 'wp:inline' in paragraph_xml),
                None,
            )
            self.assertIsNotNone(image_paragraph, 'No image paragraph found in document XML')

            self.assertIn('w:lineRule="auto"', image_paragraph)
            self.assertNotIn('w:lineRule="exact"', image_paragraph)
            self.assertIn('w:firstLine="0"', image_paragraph)
            self.assertNotIn('w:hanging=', image_paragraph)
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)

    def test_build_markdown_image_uses_caption_and_alignment_metadata(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_image_caption'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'
            output_path = workspace / 'out.docx'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Test',
                        'description': '',
                        'type': 'report',
                        'docx_template': 'template.docx',
                        'required_files': ['Ch01_Test.md'],
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: Ch01_Test.md -->\n\n'
                f'![Dang nhap]({PROJECT_ROOT.as_posix()}/test_extracted.png)'
                '{caption="Hình 4.1. Giao diện đăng nhập", width=50%, align=right}\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)
            builder.save(output_path)

            with zipfile.ZipFile(output_path, 'r') as archive:
                document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')

            self.assertIn('Hình 4.1. Giao diện đăng nhập', document_xml)
            self.assertIn('w:jc w:val="right"', document_xml)
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)

    def test_build_renders_relation_schema_code_terms_as_italic_prose(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_relation_schema'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Test',
                        'description': '',
                        'type': 'report',
                        'docx_template': 'template.docx',
                        'required_files': ['Ch01_Test.md'],
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: Ch01_Test.md -->\n\n'
                'Quan he *PROJ* duoc tach thanh `PROJ1(PNO, BUDGET)` va `PROJ2(PNO, PNAME, LOC)` voi nguong `BUDGET <= 200000`.\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)

            target_runs = {
                run.text: run
                for paragraph in builder.doc.paragraphs
                for run in paragraph.runs
                if run.text
                in {
                    'PROJ',
                    'PROJ1(PNO, BUDGET)',
                    'PROJ2(PNO, PNAME, LOC)',
                    ' voi nguong BUDGET <= 200000.',
                }
            }

            self.assertTrue(target_runs['PROJ'].italic)
            self.assertTrue(target_runs['PROJ1(PNO, BUDGET)'].italic)
            self.assertTrue(target_runs['PROJ2(PNO, PNAME, LOC)'].italic)
            self.assertFalse(bool(target_runs[' voi nguong BUDGET <= 200000.'].italic))
            self.assertNotEqual(target_runs['PROJ1(PNO, BUDGET)'].font.name, 'Courier New')
            self.assertNotEqual(target_runs['PROJ2(PNO, PNAME, LOC)'].font.name, 'Courier New')
            self.assertNotEqual(target_runs[' voi nguong BUDGET <= 200000.'].font.name, 'Courier New')
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)

    def test_build_markdown_table_cells_do_not_use_first_line_indent(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_table_indent'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Test',
                        'description': '',
                        'type': 'report',
                        'docx_template': 'template.docx',
                        'required_files': ['Ch01_Test.md'],
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: Ch01_Test.md -->\n\n| Cột 1 | Cột 2 |\n| --- | --- |\n| Nội dung 1 | Nội dung 2 |\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)

            table = builder.doc.tables[0]
            for row in table.rows:
                for cell in row.cells:
                    paragraph = cell.paragraphs[0]
                    self.assertEqual(paragraph.paragraph_format.first_line_indent.pt, 0)
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)

    def test_build_markdown_table_applies_header_and_cell_alignment_formatting(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_table_header_format'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'
            output_path = workspace / 'out.docx'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Test',
                        'description': '',
                        'type': 'report',
                        'docx_template': 'template.docx',
                        'required_files': ['Ch01_Test.md'],
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: Ch01_Test.md -->\n\n| Họ tên | Vai trò |\n| --- | --- |\n| Nguyễn Viết Tùng | Frontend |\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)
            builder.save(output_path)

            table = builder.doc.tables[0]
            header_cell = table.cell(0, 0)
            body_cell = table.cell(1, 0)

            self.assertEqual(header_cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
            self.assertEqual(body_cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
            self.assertTrue(header_cell.paragraphs[0].runs[0].bold)
            self.assertFalse(bool(body_cell.paragraphs[0].runs[0].bold))

            with zipfile.ZipFile(output_path, 'r') as archive:
                document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')

            self.assertIn('w:fill="D9EAF7"', document_xml)
            self.assertIn('w:tblHeader w:val="true"', document_xml)
            self.assertIn('<w:tcMar>', document_xml)
            self.assertIn('w:w="108" w:type="dxa"', document_xml)
            self.assertIn('w:w="60" w:type="dxa"', document_xml)
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)

    def test_default_paragraph_settings_use_times_new_roman_14(self):
        settings = DocxHelpers.get_paragraph_settings(None)

        self.assertEqual(settings['font_name'], 'Times New Roman')
        self.assertEqual(settings['font_size'], 14)

    def test_exam_cover_renders_table_layout_and_page_break(self):
        workspace = PROJECT_ROOT / 'tests' / '_tmp_docx_builder_exam_cover'
        if workspace.exists():
            shutil.rmtree(workspace, ignore_errors=True)
        workspace.mkdir(parents=True, exist_ok=True)
        try:
            config_path = workspace / 'config.yaml'
            markdown_path = workspace / 'assembled.md'
            image_cache = workspace / '.diagram_cache'
            output_path = workspace / 'out.docx'

            config_path.write_text(
                yaml.safe_dump(
                    {
                        'name': 'Exam',
                        'description': '',
                        'type': 'exam',
                        'docx_template': 'template.docx',
                        'required_files': ['F00_header.md', 'Ch01_Test.md'],
                        'settings': {
                            'cover_page': {
                                'logo_path': '',
                                'title': 'BÀI KIỂM TRA GIỮA KỲ',
                                'subject': 'MÔN: ...',
                                'lecturer_label': 'Giảng viên',
                                'lecturer_value': '...',
                                'class_label': 'Lớp',
                                'class_value': '...',
                                'student_label': 'Họ tên',
                                'student_value': '...',
                                'student_id_label': 'MSSV',
                                'student_id_value': '...',
                                'date_text': 'Hà Nội, Tháng .../....',
                            }
                        },
                    },
                    allow_unicode=True,
                    sort_keys=False,
                ),
                encoding='utf-8',
            )
            markdown_path.write_text(
                '<!-- FILE: F00_header.md -->\n\n'
                '**Tiêu đề:** BÀI KIỂM TRA GIỮA KỲ\n'
                '**Môn:** MÔN: CƠ SỞ DỮ LIỆU PHÂN TÁN\n'
                '**Giảng viên:** ThS. Mai Ngọc Lương\n'
                '**Lớp:** D24TXCN15-K\n'
                '**Họ tên:** Nguyễn Viết Tùng\n'
                '**MSSV:** K24DTCN633\n'
                '**Thời gian:** Hà Nội, Tháng 05/2026\n\n'
                '<!-- FILE: Ch01_Test.md -->\n\n'
                'Noi dung trang sau.\n',
                encoding='utf-8',
            )

            builder = DocxBuilder(workspace)
            builder.build_from_markdown(str(markdown_path), image_cache)
            builder.save(output_path)

            self.assertEqual(len(builder.doc.tables), 1)
            cover_table = builder.doc.tables[0]
            self.assertEqual(cover_table.cell(0, 0).text, 'Giảng viên')
            self.assertEqual(cover_table.cell(0, 2).text, 'ThS. Mai Ngọc Lương')
            self.assertEqual(cover_table.cell(3, 0).text, 'MSSV')
            self.assertEqual(cover_table.cell(3, 2).text, 'K24DTCN633')

            with zipfile.ZipFile(output_path, 'r') as archive:
                document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')

            self.assertIn('BÀI KIỂM TRA GIỮA KỲ', document_xml)
            self.assertIn('CƠ SỞ DỮ LIỆU PHÂN TÁN', document_xml)
            self.assertIn('Hà Nội, Tháng 05/2026', document_xml)
            self.assertIn('w:type="page"', document_xml)
            self.assertIn('Noi dung trang sau.', document_xml)
            self.assertNotIn('**Giảng viên:**', document_xml)
        finally:
            if workspace.exists():
                shutil.rmtree(workspace, ignore_errors=True)


if __name__ == '__main__':
    unittest.main()
