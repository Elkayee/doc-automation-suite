import base64
import subprocess
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document

import make


PNG_DATA = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aRX0AAAAASUVORK5CYII="
)


class MakeAndSplitTests(unittest.TestCase):
    def test_add_formatted_run_handles_nested_emphasis_without_breaking_identifiers(self) -> None:
        doc = Document()
        paragraph = doc.add_paragraph()

        make.add_formatted_run(
            paragraph,
            "_Ghi chú: **ShiftTemplate** lưu mẫu ca tái sử dụng; **Shift** là ca thực tế theo ngày; "
            "**ShiftAssignment** là kế hoạch phân công; **Attendance** ghi thực tế check_in, "
            "check_out và working_hours._",
        )

        self.assertEqual(
            paragraph.text,
            "Ghi chú: ShiftTemplate lưu mẫu ca tái sử dụng; Shift là ca thực tế theo ngày; "
            "ShiftAssignment là kế hoạch phân công; Attendance ghi thực tế check_in, "
            "check_out và working_hours.",
        )
        self.assertFalse(any("**" in run.text for run in paragraph.runs))
        self.assertTrue(any(run.text == "ShiftTemplate" and run.bold and run.italic for run in paragraph.runs))
        self.assertTrue(any(run.text == "Attendance" and run.bold and run.italic for run in paragraph.runs))

    def test_make_pipeline_uses_explicit_paths(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            chapters_dir = root / "chapters"
            chapters_dir.mkdir()
            md_out = root / "custom.md"
            docx_out = root / "custom.docx"
            cache_dir = root / "cache"

            (chapters_dir / "Ch00_header.md").write_text("# Header", encoding="utf-8")

            make.run_build_pipeline(
                chapters_dir=str(chapters_dir),
                md_out=str(md_out),
                docx_out=str(docx_out),
                img_cache=str(cache_dir),
            )

            self.assertTrue(md_out.exists())
            self.assertTrue(docx_out.exists())
            self.assertTrue(cache_dir.exists())

    def test_collect_chapter_files_includes_all_numbered_chapters_in_order(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            (root / "Ch00_header.md").write_text("header", encoding="utf-8")
            (root / "Ch00_toc.md").write_text("toc", encoding="utf-8")
            (root / "Ch01_Intro.md").write_text("intro", encoding="utf-8")
            (root / "Ch08_Foo.md").write_text("foo", encoding="utf-8")
            (root / "Ch10_Bar.md").write_text("bar", encoding="utf-8")

            files = [Path(path).name for path in make.collect_chapter_files(str(root))]

            self.assertEqual(
                files,
                ["Ch00_header.md", "Ch00_toc.md", "Ch01_Intro.md", "Ch08_Foo.md", "Ch10_Bar.md"],
            )

    def test_step_assemble_merges_chapters_and_skips_ch08_ch09(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            chapters_dir = root / "chapters"
            output_path = root / "report.md"
            chapters_dir.mkdir()

            (chapters_dir / "Ch00_Header.md").write_text("Header", encoding="utf-8")
            (chapters_dir / "Ch01_Intro.md").write_text("Intro", encoding="utf-8")
            (chapters_dir / "Ch08_Keep.md").write_text("Keep 8", encoding="utf-8")
            (chapters_dir / "Ch10_Keep.md").write_text("Keep 10", encoding="utf-8")

            with mock.patch.object(make, "CH_DIR", str(chapters_dir)), mock.patch.object(
                make, "MD_OUT", str(output_path)
            ):
                make.step_assemble()

            content = output_path.read_text(encoding="utf-8")
            self.assertEqual(content, "Header\n\nIntro\n\nKeep 8\n\nKeep 10")

    def test_split_script_creates_numbered_chapter_files(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source_path = root / "Bao_Cao_Tieu_Luan_NMCNPM.md"
            source_path.write_text(
                "# Report\n\nIntro block\n\n## Chapter One\nA\n\n## Chapter: Two\nB\n",
                encoding="utf-8",
            )

            completed = subprocess.run(
                ["python", str(Path.cwd() / "split_chapters.py")],
                cwd=root,
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
            )

            chapter_files = sorted((root / "chapters").glob("*.md"))
            self.assertEqual(
                [path.name for path in chapter_files],
                ["Ch00_header.md", "Ch01_Chapter_One.md", "Ch02_Chapter_Two.md"],
            )
            self.assertIn("Da tao 3 file chapter", completed.stdout)

    def test_parse_and_write_supports_centered_text_and_markdown_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            image_path = root / "logo.png"
            image_path.write_bytes(PNG_DATA)

            markdown_path = root / "cover.md"
            markdown_path.write_text(
                f"<center>**BÁO CÁO BÀI TẬP LỚN**</center>\n\n![PTIT]({image_path.as_posix()})\n",
                encoding="utf-8",
            )

            doc = Document()
            make.parse_and_write(doc, str(markdown_path))

            rendered_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            self.assertEqual(rendered_paragraphs[0].text.strip(), "BÁO CÁO BÀI TẬP LỚN")
            self.assertEqual(rendered_paragraphs[0].alignment, make.WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(len(doc.inline_shapes), 1)

    def test_parse_and_write_preserves_italic_note_with_bold_terms_and_identifiers(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown_path = root / "note.md"
            markdown_path.write_text(
                "_Ghi chú: **ShiftTemplate** lưu mẫu ca tái sử dụng; **Shift** là ca thực tế theo ngày; "
                "**ShiftAssignment** là kế hoạch phân công; **Attendance** ghi thực tế check_in, "
                "check_out và working_hours._\n",
                encoding="utf-8",
            )

            doc = Document()
            make.parse_and_write(doc, str(markdown_path))

            paragraph = next(p for p in doc.paragraphs if "ShiftTemplate" in p.text)
            self.assertEqual(
                paragraph.text,
                "Ghi chú: ShiftTemplate lưu mẫu ca tái sử dụng; Shift là ca thực tế theo ngày; "
                "ShiftAssignment là kế hoạch phân công; Attendance ghi thực tế check_in, "
                "check_out và working_hours.",
            )
            self.assertTrue(any(run.text == "ShiftTemplate" and run.bold and run.italic for run in paragraph.runs))
            self.assertTrue(any("working_hours." in run.text and run.italic for run in paragraph.runs))

    def test_markdown_to_preview_text_removes_raw_markdown_markers_in_fallback(self) -> None:
        markdown_content = (
            "## CHƯƠNG 2: Xây dựng Biểu đồ lớp thực thể\n\n"
            "Các lớp cốt lõi theo nguyên tắc **Separation of Concerns**.\n\n"
            "| **Tên thực thể** | **Mô tả** |\n"
            "| --- | --- |\n"
            "| Store | Cửa hàng |\n"
            "| Order | Đơn hàng |\n"
        )

        preview_text = make.markdown_to_preview_text(markdown_content)

        self.assertIn("CHƯƠNG 2: Xây dựng Biểu đồ lớp thực thể", preview_text)
        self.assertIn("Separation of Concerns", preview_text)
        self.assertIn("Tên thực thể | Mô tả", preview_text)
        self.assertIn("Store | Cửa hàng", preview_text)
        self.assertNotIn("##", preview_text)
        self.assertNotIn("**", preview_text)
        self.assertNotIn("| --- | --- |", preview_text)

    def test_markdown_to_preview_blocks_parses_heading_quote_and_table_for_fallback_renderer(self) -> None:
        markdown_content = (
            "## CHƯƠNG 1\n\n"
            "> **Mục tiêu** chương này.\n\n"
            "Các lớp cốt lõi theo nguyên tắc **Separation of Concerns**.\n\n"
            "| **Tên thực thể** | **Mô tả** |\n"
            "| --- | --- |\n"
            "| Store | Cửa hàng |\n"
        )

        blocks = make.markdown_to_preview_blocks(markdown_content)

        self.assertEqual(blocks[0]["type"], "heading")
        self.assertEqual(blocks[0]["level"], 2)
        self.assertEqual("".join(text for text, _ in blocks[0]["segments"]), "CHƯƠNG 1")

        self.assertEqual(blocks[1]["type"], "quote")
        self.assertTrue(any(text == "Mục tiêu" and "bold" in styles for text, styles in blocks[1]["segments"]))

        self.assertEqual(blocks[2]["type"], "paragraph")
        self.assertTrue(
            any(text == "Separation of Concerns" and "bold" in styles for text, styles in blocks[2]["segments"])
        )

        self.assertEqual(blocks[3]["type"], "table")
        self.assertEqual(blocks[3]["rows"][0], ["Tên thực thể", "Mô tả"])
        self.assertEqual(blocks[3]["rows"][1], ["Store", "Cửa hàng"])

    def test_parse_and_write_resolves_repo_relative_image_paths(self) -> None:
        relative_asset = Path("extracted_media/test_ui_logo.png")
        asset_path = Path(make.BASE) / relative_asset
        asset_path.parent.mkdir(parents=True, exist_ok=True)
        asset_path.write_bytes(PNG_DATA)

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                root = Path(temp_dir)
                markdown_path = root / "cover.md"
                markdown_path.write_text(
                    "![PTIT](../extracted_media/test_ui_logo.png)\n",
                    encoding="utf-8",
                )

                doc = Document()
                make.parse_and_write(doc, str(markdown_path))
                self.assertEqual(len(doc.inline_shapes), 1)
        finally:
            if asset_path.exists():
                asset_path.unlink()

    def test_parse_and_write_inserts_word_toc_and_indents_numbered_headings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            markdown_path = root / "toc.md"
            markdown_path.write_text(
                "## MỤC LỤC\n[[TOC]]\n\n## 2.1 Xác định các ACTOR\n",
                encoding="utf-8",
            )

            doc = Document()
            make.parse_and_write(doc, str(markdown_path))

            heading_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            self.assertEqual(heading_paragraphs[0].text.strip(), "MỤC LỤC")
            self.assertIn('TOC \\o "1-4" \\h \\z \\u', doc.paragraphs[1]._p.xml)
            self.assertIn('<w:r>', doc.paragraphs[1]._p.xml)
            self.assertIn('w:fldCharType="begin"', doc.paragraphs[1]._p.xml)
            self.assertEqual(heading_paragraphs[-1].text.strip(), "2.1 Xác định các ACTOR")
            self.assertAlmostEqual(heading_paragraphs[-1].paragraph_format.left_indent.cm, 0.5, places=1)


if __name__ == "__main__":
    unittest.main()
