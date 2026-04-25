import base64
import tempfile
import unittest
from pathlib import Path

from docx import Document

from convert_docx_to_md import convert_docx_to_markdown


PNG_DATA = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aRX0AAAAASUVORK5CYII="
)


class ConvertDocxToMarkdownTests(unittest.TestCase):
    def test_converts_headings_paragraphs_tables_and_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            image_path = root / "sample.png"
            image_path.write_bytes(PNG_DATA)

            input_path = root / "sample.docx"
            output_path = root / "sample.md"
            media_dir = root / "media"

            doc = Document()
            doc.add_heading("Test Report", level=1)

            paragraph = doc.add_paragraph()
            paragraph.add_run("Bold").bold = True
            paragraph.add_run(" and ")
            paragraph.add_run("italic").italic = True

            image_paragraph = doc.add_paragraph()
            image_paragraph.add_run().add_picture(str(image_path))

            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Alpha"
            table.cell(1, 1).text = "42"

            doc.save(input_path)

            result = convert_docx_to_markdown(input_path, output_path, media_dir)

            markdown = output_path.read_text(encoding="utf-8")

            self.assertEqual(result, output_path)
            self.assertIn("# Test Report", markdown)
            self.assertIn("**Bold** and *italic*", markdown)
            self.assertIn("| Name | Value |", markdown)
            self.assertIn("| --- | --- |", markdown)
            self.assertIn("![Image 1](", markdown)
            self.assertTrue(any(media_dir.iterdir()))


if __name__ == "__main__":
    unittest.main()
